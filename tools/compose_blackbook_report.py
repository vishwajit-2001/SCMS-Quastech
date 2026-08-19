from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
BASE_DOC = Path(r"D:\Downlaods\Inital Pages of Black Book.docx")
SOURCE_REPORT = ROOT / "scms_sem_iv_project_report.docx"
OUTPUT = ROOT / "scms_sem_iv_project_report_blackbook.docx"
JAVA_ROOT = ROOT / "src" / "main" / "java" / "com" / "scms" / "scms"
TEMPLATE_ROOT = ROOT / "src" / "main" / "resources" / "templates"

TITLE = "Smart Campus Management System (SCMS)"
STUDENT_NAMES = "Vishwajit Gadale, Amar Bharati, Aakash Vishwakarma"


def set_page_borders(section) -> None:
    sectPr = section._sectPr
    pgBorders = sectPr.find(qn("w:pgBorders"))
    if pgBorders is None:
        pgBorders = OxmlElement("w:pgBorders")
        sectPr.append(pgBorders)
    for edge in ["top", "left", "bottom", "right"]:
        el = pgBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            pgBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "2")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def replace_placeholders(doc: Document) -> None:
    for para in doc.paragraphs:
        full = para.text
        updated = (
            full.replace("<Title>", TITLE)
            .replace("“Title”", f"“{TITLE}”")
            .replace('"Title"', f'"{TITLE}"')
            .replace("<Student Name>", STUDENT_NAMES)
            .replace("<Name>", STUDENT_NAMES)
        )
        if updated != full:
            if para.runs:
                para.runs[0].text = updated
                for run in para.runs[1:]:
                    run.text = ""


def find_body_start(doc: Document) -> object:
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.iter(qn("w:t")))
            if text.strip() == "Table of Contents":
                return child
    raise RuntimeError("Could not find 'Table of Contents' in source report.")


def append_from_marker(base: Document, source: Document) -> int:
    body = source.element.body
    start = find_body_start(source)
    started = False
    appended = 0
    for child in body.iterchildren():
        if child is start:
            started = True
        if not started:
            continue
        if child.tag == qn("w:sectPr"):
            continue
        base.element.body.append(deepcopy(child))
        appended += 1
    return appended


def apply_template_styles(doc: Document, paragraph_start: int) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    if "Default" in [s.name for s in styles]:
        styles["Default"].font.name = "Times New Roman"
        styles["Default"].font.size = Pt(12)

    for style_name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True

    for para in doc.paragraphs[paragraph_start:]:
        text = para.text.strip()
        if para.style and para.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif text == "Table of Contents":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(18)
                run.bold = True
        elif any(run.element.xpath(".//w:drawing") for run in para.runs):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text.startswith("Figure "):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
                run.italic = True
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in para.runs:
            if para.style and para.style.name in {"Heading 1", "Heading 2", "Heading 3"}:
                run.font.name = "Times New Roman"
            elif text == "Table of Contents":
                continue
            elif text.startswith("Figure "):
                continue
            else:
                if run.font.bold is None:
                    run.font.bold = False
                run.font.name = "Times New Roman"
                if run.font.size is None:
                    run.font.size = Pt(12)

    for table in doc.tables:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        if run.font.size is None:
                            run.font.size = Pt(11)
                        if row_index == 0:
                            run.font.bold = True


def clear_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        for para in footer.paragraphs:
            para.text = ""


def pretty_name(name: str) -> str:
    return name.replace(".java", "").replace(".html", "").replace("-", " ").replace("_", " ")


def add_appendices(doc: Document) -> None:
    controllers = sorted((JAVA_ROOT / "controller").glob("*.java"))
    models = sorted((JAVA_ROOT / "model").glob("*.java"))
    screens = sorted(TEMPLATE_ROOT.rglob("*.html"))

    doc.add_page_break()
    doc.add_heading("Appendix A: Detailed Controller Catalog", level=1)
    doc.add_paragraph(
        "This appendix documents the controller-level implementation in additional detail so that the report reflects the actual code structure behind the campus platform."
    )
    for path in controllers:
        name = pretty_name(path.name)
        doc.add_heading(name, level=2)
        doc.add_paragraph(
            f"{name} is part of the request-handling layer of the SCMS application. It groups HTTP endpoints that belong to one operational area and prepares model data for the corresponding Thymeleaf templates or backend action flows."
        )
        doc.add_paragraph(
            f"From a maintainability perspective, {name} helps keep the campus management system modular. It reduces coupling between unrelated user journeys and makes role-aware access, validation, and view routing easier to understand during development, debugging, and future enhancement work."
        )

    doc.add_page_break()
    doc.add_heading("Appendix B: Detailed Entity Catalog", level=1)
    doc.add_paragraph(
        "The entity catalog below expands the database and domain discussion by describing the purpose of the classes present in the model package."
    )
    for path in models:
        name = pretty_name(path.name)
        doc.add_heading(name, level=2)
        doc.add_paragraph(
            f"{name} represents a structured domain record within SCMS. It exists so that business workflows can operate on strongly modeled institutional data instead of disconnected manual entries."
        )
        doc.add_paragraph(
            f"In practical terms, {name} contributes to traceability, persistence consistency, and relationship-aware processing. The class is part of the broader object model that enables academic, administrative, notification, assessment, and placement features to share a common database-backed language."
        )

    doc.add_page_break()
    doc.add_heading("Appendix C: Detailed Screen Inventory", level=1)
    doc.add_paragraph(
        "This appendix lists the implemented HTML templates and explains their role in the overall portal surface. It complements the earlier I/O screen layout section with a fuller inventory view."
    )
    for path in screens:
        rel = path.relative_to(TEMPLATE_ROOT)
        name = pretty_name(path.name).title()
        role = rel.parts[0]
        doc.add_heading(name, level=2)
        doc.add_paragraph(
            f"The template `{rel.as_posix()}` belongs to the {role} area of the application. It is designed to support a focused workflow screen within that portal and helps keep the user experience aligned with the responsibilities of the current role."
        )
        doc.add_paragraph(
            f"By separating `{name}` into its own template, the project keeps navigation, form behavior, status display, and content rendering manageable. This improves front-end maintainability and makes future refinements easier without disturbing unrelated screens."
        )


def main() -> None:
    base = Document(str(BASE_DOC))
    source = Document(str(SOURCE_REPORT))

    initial_paragraphs = len(base.paragraphs)
    replace_placeholders(base)
    append_from_marker(base, source)
    add_appendices(base)
    apply_template_styles(base, initial_paragraphs)
    clear_footer(base)

    sec = base.sections[0]
    set_page_borders(sec)
    base.save(str(OUTPUT))


if __name__ == "__main__":
    main()
