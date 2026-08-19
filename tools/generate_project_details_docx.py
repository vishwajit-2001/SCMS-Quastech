from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


OUTPUT_PATH = r"D:\internship\scms-project\scms-project\updated_project_details_v3.docx"


def set_font(run, size, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, 18, bold=True, color="12355B")
    p.space_after = Pt(4)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, 10, color="465A69")
    p.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_font(run, 13, bold=True, color="12355B")
        p.space_before = Pt(6)
        p.space_after = Pt(3)
    else:
        set_font(run, 10.5, bold=True, color="214E6E")
        p.space_before = Pt(4)
        p.space_after = Pt(2)


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, 9.5)
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.15


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_font(run, 9.5)
        pf = p.paragraph_format
        pf.space_after = Pt(1)
        pf.line_spacing = 1.1


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p.add_run("Page ")
    r = p.add_run()
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_end)


doc = Document()
section = doc.sections[0]
section.top_margin = Pt(42)
section.bottom_margin = Pt(42)
section.left_margin = Pt(52)
section.right_margin = Pt(52)
add_page_number(section)

add_title(doc, "AI-Powered Smart Campus Management System")
add_subtitle(doc, "Updated Project Details Based on Current Implementation")
add_subtitle(doc, "Project: SCMS")
add_subtitle(doc, "Updated on: 23 May 2026")

add_heading(doc, "Group Members (Group A)", 2)
add_bullets(doc, ["Vishwajit Gadale", "Amar Bharati", "Aakash Vishwakarma"])

add_heading(doc, "Project Overview")
add_body(doc, "SCMS is a centralized web-based campus management platform for multi-college academic operations. The current implementation is a monolithic Spring Boot application with separate role-based portals for Super Admin, College Admin, Teacher, Student, and Placement TPO users.")
add_body(doc, "The system manages student onboarding, academic structure, classroom operations, attendance, assignments, study materials, fees, exams, events, placement workflows, notifications, and tenant-level administration from a single application.")

add_heading(doc, "Key Feature Modules")
add_heading(doc, "Student Module", 2)
add_bullets(doc, [
    "College-wise registration and login",
    "Detailed onboarding workflow with personal, academic, address, and document sections",
    "Timetable, subjects, assignments, study materials, fees, placements, events, and notifications",
    "Assignment submission and downloadable academic resources",
    "Exam and quiz participation with result visibility",
])
add_heading(doc, "TPO Module", 2)
add_bullets(doc, [
    "Placement TPO login with dedicated dashboard and college-linked workspace",
    "Company drive creation and placement-cycle management",
    "Student application review, interview scheduling, and analytics tracking",
    "Excel export of placement applications and ZIP export of submitted resumes",
    "Placement reports and company-wise workflow visibility",
])
add_heading(doc, "Teacher Module", 2)
add_bullets(doc, [
    "Teacher dashboard with assigned students and timetable visibility",
    "Attendance session creation and status updates",
    "Assignment publishing with file upload and submission review",
    "Study material upload and class resource management",
    "Standard exam scheduling and PDF-based auto-generated quiz creation",
])
add_heading(doc, "Admin Module", 2)
add_bullets(doc, [
    "Student, teacher, course, batch, class, subject, and academic structure management",
    "Fee setup and operational reports",
    "Placement drive, TPO, application, and interview management",
    "Import/export support for operational data and spreadsheets",
    "Portal notifications and document review handling",
])
add_heading(doc, "Super Admin and Platform Services", 2)
add_bullets(doc, [
    "College lifecycle management with active, inactive, and setup-pending states",
    "College-admin provisioning, communication center, notifications, and audit logs",
    "College-wise portal isolation using college codes",
    "Role-aware chatbot support for student and teacher self-service queries",
    "Sensitive field encryption support for Aadhaar, bank, and related data",
])

add_heading(doc, "Updated Tech Architecture")
add_bullets(doc, [
    "Backend framework: Spring Boot 3.2.3 on Java 17",
    "Server-side web layer: Spring MVC with Thymeleaf templates",
    "Persistence: Spring Data JPA with MySQL",
    "Security model: BCrypt password hashing with session-based role access",
    "Utilities: Apache POI for Excel export/import, PDFBox for PDF text extraction, JavaMail for email support",
    "Frontend delivery: HTML, CSS, JavaScript, Thymeleaf-rendered portals",
    "Static file handling for photos, signatures, resumes, exam files, assignments, and student documents",
])

add_heading(doc, "AI Modules Explanation")
add_body(doc, "The current SCMS implementation includes practical AI-like and automation-oriented modules, but it is not yet a full machine-learning platform. The project currently uses intelligent workflow logic, automated content extraction, and contextual assistance features that support academic operations.")
add_heading(doc, "Implemented AI-Related Modules", 2)
add_bullets(doc, [
    "Chatbot assistance module for student and teacher queries based on role, profile, timetable, fees, assignments, subjects, and attendance context.",
    "PDF-based quiz generation workflow in the exam automation service, where uploaded study PDFs are parsed and converted into quiz-ready question data.",
    "Automation-oriented exam support for scheduling, notifications, attempt tracking, and response handling.",
    "Rule-based academic and portal assistance that improves self-service access without requiring manual staff intervention for common questions.",
])
add_heading(doc, "Current AI Scope Clarification", 2)
add_bullets(doc, [
    "The current version does not implement a separate ML model training pipeline.",
    "Predictive student-risk analytics and advanced recommendation engines are not part of the implemented codebase yet.",
    "The present solution should be described as AI-assisted or automation-enabled rather than as a full predictive AI platform.",
])

add_heading(doc, "API Modules")
add_body(doc, "SCMS is primarily built as a Spring Boot web application with controller-driven modules. Even though many flows render Thymeleaf pages directly, the application is still organized around clear backend modules and request endpoints.")
add_heading(doc, "Core API / Controller Modules", 2)
add_bullets(doc, [
    "MainController: registration, login, and portal entry management.",
    "AdminController: student, teacher, class, timetable, fee, subject, course, batch, placement, and report operations.",
    "StudentController: onboarding, dashboard, assignment submission, fees, timetable, exams, placements, and profile workflows.",
    "TeacherController: dashboard, attendance, assignments, study materials, timetable, and exam operations.",
    "PlacementController: TPO login, company drives, applications, interviews, analytics, reports, and export operations.",
    "SuperAdminController: college management, admin management, communications, audit logs, notifications, and tenant oversight.",
    "EventController: campus event creation and participation workflows.",
    "ChatbotController: chatbot context loading and message handling endpoints.",
])
add_heading(doc, "Service Modules Supporting APIs", 2)
add_bullets(doc, [
    "ExamAutomationService for exam creation, quiz generation, notifications, and attempt handling.",
    "ChatbotService for contextual response generation.",
    "AcademicStructureService for course, semester, and subject mapping support.",
    "PasswordProtectionService for password hashing and matching.",
    "EmployeeIdService and TeacherAcademicMappingService for operational identity and academic assignments.",
])

add_heading(doc, "BRS - Business Requirement Specification")
add_heading(doc, "1.1 Purpose", 2)
add_body(doc, "To provide a centralized campus platform that reduces manual coordination across colleges and improves student, academic, and administrative visibility through one integrated web system.")
add_heading(doc, "1.2 Business Objectives", 2)
add_bullets(doc, [
    "Digitize student admission-to-portal onboarding workflows",
    "Reduce operational work for admins, teachers, and placement teams",
    "Standardize class, course, batch, and subject management",
    "Improve communication using in-app notifications and college-level portals",
    "Support multi-college oversight through a super admin control layer",
])
add_heading(doc, "1.3 Stakeholders", 2)
add_bullets(doc, [
    "Super Admin",
    "College Management",
    "College Admin Staff",
    "Teachers",
    "Students",
    "Placement TPO",
])
add_heading(doc, "1.4 Business Scope", 2)
add_heading(doc, "In Scope", 2)
add_bullets(doc, [
    "College onboarding and status control",
    "Student, teacher, TPO, course, batch, subject, and classroom management",
    "Attendance, assignments, study materials, exams, fees, events, and placements",
    "Student document collection and onboarding verification",
    "Notifications, chatbot assistance, audit logs, and operational exports",
])
add_heading(doc, "Out of Scope", 2)
add_bullets(doc, [
    "Biometric device integration",
    "Native mobile applications",
    "External government or university API integrations",
    "Dedicated microservice deployment in the current version",
])

add_heading(doc, "SRS - Software Requirement Specification")
add_heading(doc, "2.1 Introduction", 2)
add_body(doc, "This document captures the current implemented scope and aligned requirements for the SCMS platform.")
add_heading(doc, "2.2 Overall System Description", 2)
add_bullets(doc, [
    "Product type: browser-based enterprise web application",
    "Deployment style: single Spring Boot application",
    "User classes: Super Admin, Admin, Teacher, Student, Placement TPO",
    "Access model: college-code-aware portal entry with session-backed role routing",
])
add_heading(doc, "2.3 System Features (High Level)", 2)
add_bullets(doc, [
    "Authentication, password protection, and role-aware session access",
    "Student onboarding and student lifecycle management",
    "Academic structure and timetable operations",
    "Assignment, material, exam, and quiz workflows",
    "Fee, placement, event, notification, and reporting modules",
    "Super admin oversight and chatbot support",
])

add_heading(doc, "Functional Requirements (FRS)")
add_heading(doc, "3.1 User and Access Management", 2)
add_bullets(doc, [
    "FR1: System shall allow student self-registration through college-specific registration links.",
    "FR2: System shall support login for Admin, Teacher, Student, Super Admin, and Placement TPO users.",
    "FR3: System shall maintain session-based access control for protected routes.",
    "FR4: System shall store passwords using BCrypt hashing and support legacy password migration.",
])
add_heading(doc, "3.2 College and Administration Management", 2)
add_bullets(doc, [
    "FR5: Super Admin shall manage colleges and college activation status.",
    "FR6: Super Admin shall manage college-admin communication and audit logs.",
    "FR7: Admin shall manage courses, batches, subjects, classrooms, and academic structure.",
    "FR8: Admin shall manage students, teachers, and placement TPO accounts.",
])
add_heading(doc, "3.3 Student Lifecycle and Onboarding", 2)
add_bullets(doc, [
    "FR9: Students shall complete a multi-section onboarding form after registration.",
    "FR10: System shall store student profile, address, academic, and uploaded document data.",
    "FR11: Admin shall review or request re-upload of student documents.",
    "FR12: Students shall view profile, subjects, timetable, and notifications from their portal.",
])
add_heading(doc, "3.4 Teaching and Academic Operations", 2)
add_bullets(doc, [
    "FR13: Teachers shall mark attendance for timetable-linked sessions.",
    "FR14: Teachers shall upload assignments and study materials.",
    "FR15: Students shall submit assignments and access study materials.",
    "FR16: System shall maintain class-wise timetable visibility for students and teachers.",
])
add_heading(doc, "3.5 Exam and Quiz Management", 2)
add_bullets(doc, [
    "FR17: Teachers shall schedule standard exams with instructions, topics, files, and room details.",
    "FR18: System shall create quiz exams by extracting text from uploaded PDF notes.",
    "FR19: System shall store quiz questions, attempts, and responses.",
    "FR20: Students shall view upcoming exams, attempt quizzes, and review results where available.",
])
add_heading(doc, "3.6 Fees, Placements, and Events", 2)
add_bullets(doc, [
    "FR21: Admin shall define fee records and students shall view fee summaries.",
    "FR22: Placement TPO shall manage drives, applications, interviews, analytics, and reports.",
    "FR23: System shall export placement applications to Excel and selected resumes to ZIP.",
    "FR24: Students shall access event and placement information from the portal.",
])
add_heading(doc, "3.7 Notifications and Assistance", 2)
add_bullets(doc, [
    "FR25: System shall support in-app notifications across roles.",
    "FR26: Super Admin shall send messages to college admins.",
    "FR27: Chatbot shall answer role-aware questions for timetable, fees, assignments, and profile context.",
])

add_heading(doc, "Non-Functional Requirements (NFR)")
add_heading(doc, "4.1 Performance", 2)
add_bullets(doc, [
    "NFR1: Common dashboard and list views should load within reasonable web response times under normal campus usage.",
    "NFR2: File upload and export operations shall handle moderate academic-office workloads.",
])
add_heading(doc, "4.2 Security", 2)
add_bullets(doc, [
    "NFR3: Passwords must be encrypted using BCrypt.",
    "NFR4: Sensitive fields shall support encryption using the configured application secret.",
    "NFR5: Protected routes shall enforce session-aware role validation.",
    "NFR6: File access paths shall be normalized to avoid unsafe path traversal.",
])
add_heading(doc, "4.3 Maintainability", 2)
add_bullets(doc, [
    "NFR7: System shall remain modular through controllers, services, repositories, and model entities.",
    "NFR8: Schema evolution shall be supportable through runner-based migration utilities and normalized entities.",
    "NFR9: UI templates shall remain separated by role for maintainable portal updates.",
])
add_heading(doc, "4.4 Reliability and Operations", 2)
add_bullets(doc, [
    "NFR10: Uploaded academic files and portal assets shall remain accessible through managed static storage.",
    "NFR11: Audit and notification records shall preserve administrative activity history.",
    "NFR12: The application shall support development-time hot reload for faster maintenance cycles.",
])

add_heading(doc, "Technical Requirements")
add_heading(doc, "Backend", 2)
add_bullets(doc, ["Java 17", "Spring Boot", "Spring MVC", "Spring Data JPA / Hibernate", "Spring Security", "Maven"])
add_heading(doc, "Database", 2)
add_bullets(doc, [
    "MySQL",
    "Normalized relational schema",
    "Entity relationships for colleges, admins, teachers, students, academics, exams, and placements",
])
add_heading(doc, "Frontend", 2)
add_bullets(doc, ["Thymeleaf templates", "HTML/CSS/JavaScript", "Role-specific dashboards and portal screens"])
add_heading(doc, "Document and File Utilities", 2)
add_bullets(doc, ["Apache POI", "Apache PDFBox", "Multipart file handling for uploads and exports"])

add_heading(doc, "High-Level Architecture")
add_body(doc, "The current system follows a single-application modular architecture. Role-specific controllers handle web requests, services implement core business logic, repositories manage persistence, and Thymeleaf templates deliver the role portals.")
add_bullets(doc, [
    "Presentation layer: Thymeleaf pages, static CSS, and JavaScript",
    "Controller layer: Main, Admin, Student, Teacher, Placement, Super Admin, Event, and Chatbot controllers",
    "Service layer: academic structure, password protection, employee IDs, chatbot, and exam automation",
    "Persistence layer: JPA repositories and MySQL-backed entities",
    "Storage layer: local static upload directories for academic and profile documents",
])

add_heading(doc, "Core Database Tables / Entities")
add_bullets(doc, [
    "college",
    "admin",
    "student",
    "teacher",
    "placement_tpo",
    "course",
    "batch",
    "academic_structure",
    "subject and subject_master",
    "classroom and timetable",
    "teacher_attendance_session and teacher_attendance_entry",
    "assignment and assignment_submission",
    "study_material",
    "fees",
    "exam_session, exam_attendance, exam_quiz_question, exam_quiz_attempt, exam_quiz_response",
    "campus_event and campus_event_application",
    "placement_drive, placement_application, placement_interview",
    "portal_notification and super_admin_audit_log",
    "student_document",
])

add_heading(doc, "Database Design")
add_body(doc, "The SCMS database follows a normalized relational structure in MySQL. The design separates platform administration, academic structure, user accounts, academic operations, notifications, and placement workflows into dedicated entities with mapped relationships.")
add_heading(doc, "Database Design Principles", 2)
add_bullets(doc, [
    "College-aware tenancy through college, admin, and role-linked user records.",
    "Separate master data for course, batch, academic structure, subject, and classroom mapping.",
    "Operational entities for attendance, assignments, exams, fees, placements, notifications, and events.",
    "Document and file metadata stored separately from actual uploaded files.",
    "Extensible schema supported through migration runner classes and incremental schema evolution.",
])
add_heading(doc, "Major Relationship Flows", 2)
add_bullets(doc, [
    "A college is linked to one or more admin accounts, and each admin manages institution-level academic data.",
    "Students, teachers, and placement TPO users are linked to the owning admin or college context.",
    "Courses, batches, subjects, classrooms, and timetables define the academic structure visible to users.",
    "Attendance, assignments, study materials, fees, exams, and placements are mapped to the relevant student, teacher, admin, or classroom context.",
    "Notifications and audit logs capture cross-module communication and traceability.",
])

add_heading(doc, "UI/UX Design Requirements")
add_body(doc, "The SCMS user interface should remain practical, role-oriented, and easy to navigate for daily campus operations. The implemented Thymeleaf portal structure already follows role-based separation, and the design requirements below capture the intended UI/UX standard for the project.")
add_heading(doc, "Functional UI/UX Requirements", 2)
add_bullets(doc, [
    "Separate dashboard experiences for Super Admin, Admin, Teacher, Student, and Placement TPO users.",
    "College-specific portal branding and login flow using college codes and college-aware registration links.",
    "Clear navigation for recurring workflows such as onboarding, attendance, assignments, fees, exams, and placements.",
    "Readable tables, forms, status labels, notifications, and file-upload interfaces for daily operational use.",
    "Responsive layout behavior so users can access core workflows from standard desktop and laptop browsers.",
])
add_heading(doc, "Design Quality Requirements", 2)
add_bullets(doc, [
    "Interfaces should minimize confusion by keeping module names and actions consistent across roles.",
    "Forms should support large academic datasets without becoming cluttered.",
    "Student onboarding should be segmented into manageable steps instead of one long unstructured form.",
    "Dashboards should emphasize quick scanning of counts, statuses, pending work, and announcements.",
    "Upload, download, and export actions should be clearly visible and easy to confirm.",
])

add_heading(doc, "Future Enhancements")
add_body(doc, "The current implementation is functionally broad, but several enhancements can strengthen the project further in future iterations.")
add_heading(doc, "Recommended Next Enhancements", 2)
add_bullets(doc, [
    "Advanced predictive analytics for attendance risk, dropout signals, and academic performance trends.",
    "Dedicated REST API standardization for all modules to support mobile apps or third-party integrations.",
    "Role-based audit enrichment and stronger production-grade authorization hardening.",
    "Online fee payment gateway integration and receipt-generation workflow improvements.",
    "Parent portal or guardian visibility module if required by the institution.",
    "Real-time notifications using WebSocket or push-based communication.",
    "Dashboard analytics with richer charts and longitudinal academic comparisons.",
    "Cloud file storage, backups, and deployment automation for production readiness.",
    "Swagger or OpenAPI documentation for maintainable API exposure.",
    "Mobile-friendly or native app support for student and teacher self-service.",
])

add_heading(doc, "Updated Team Division")
add_bullets(doc, [
    "Member 1 - Core backend, authentication flows, repositories, and security-related logic",
    "Member 2 - Academic operations, exams, fees, placements, and admin workflows",
    "Member 3 - Thymeleaf UI, dashboards, role portals, and integration/testing support",
])
add_body(doc, "This revised project details document reflects the current implemented SCMS project rather than the earlier planned stack.")

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
