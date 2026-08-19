from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "SCMS_Interview_Preparation_Guide.docx"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
        shade(table.rows[0].cells[idx], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def qa(doc, question, answer):
    p = doc.add_paragraph()
    r = p.add_run("Q. " + question)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    doc.add_paragraph("A. " + answer)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
for style_name, size, color in [
    ("Heading 1", 16, RGBColor(46, 116, 181)),
    ("Heading 2", 13, RGBColor(46, 116, 181)),
    ("Heading 3", 12, RGBColor(31, 77, 120)),
]:
    styles[style_name].font.name = "Calibri"
    styles[style_name].font.size = Pt(size)
    styles[style_name].font.color.rgb = color

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SCMS Project Interview Preparation Guide")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(11, 37, 69)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Student College Management System | Technical Architecture, Database, APIs, and Interview Q&A").italic = True
doc.add_paragraph()

heading(doc, "1. Project Overview")
doc.add_paragraph(
    "SCMS is a Spring Boot based Student College Management System for managing college operations across multiple roles: "
    "super admin, institute admin, students, teachers, and placement/TPO users. The system supports admission/onboarding, "
    "student and teacher management, academic structure, courses, batches, subjects, timetables, fees, attendance, assignments, "
    "exams/quizzes, study materials, campus events, placements, notifications, reports, and document verification."
)
bullet(doc, "Architecture style: server-side MVC web application with selected JSON endpoints.")
bullet(doc, "Backend: Java 17 and Spring Boot 3.2.3.")
bullet(doc, "Frontend: Thymeleaf templates, HTML, CSS, JavaScript, and static assets.")
bullet(doc, "Persistence: MySQL relational database accessed through Spring Data JPA/Hibernate.")
bullet(doc, "Packaging: Maven JAR application.")

heading(doc, "2. Technology Stack")
add_table(doc, ["Layer", "Technology", "Purpose"], [
    ("Language", "Java 17", "Primary backend language with modern LTS support."),
    ("Framework", "Spring Boot 3.2.3", "Application bootstrapping, dependency management, MVC, and configuration."),
    ("Web", "Spring MVC", "Controllers, request mappings, form handling, file download/upload endpoints."),
    ("View Engine", "Thymeleaf + thymeleaf-extras-springsecurity6", "Server-rendered pages and role-based template support."),
    ("Database", "MySQL", "Relational persistence for normalized academic and operational data."),
    ("ORM", "Spring Data JPA + Hibernate", "Entity mapping, repositories, query derivation, JPQL queries."),
    ("Security", "Spring Security + BCrypt + custom interceptor", "Password hashing and session/role access checks."),
    ("File Processing", "Multipart upload, Apache POI, PDFBox", "Student documents, resumes, assignments, Excel/PDF exports."),
    ("Frontend Assets", "HTML/CSS/JavaScript/SVG", "Dashboard UI, onboarding screens, chatbot widget, portal layouts."),
    ("Build Tool", "Maven", "Dependency management, compile, test, and Spring Boot packaging."),
])

heading(doc, "3. Database Type and Design")
doc.add_paragraph(
    "The project uses a relational database, not a NoSQL/non-relational database. This is appropriate because the domain has "
    "strong relationships and integrity requirements: admins own courses, courses have batches, batches have students and subjects, "
    "teachers are assigned to academic structures, placement drives receive applications, and interviews select students."
)
bullet(doc, "Database engine: MySQL using the mysql-connector-j runtime driver.")
bullet(doc, "Configured datasource: jdbc:mysql://localhost:3306/scms in application.properties.")
bullet(doc, "Schema management: Hibernate ddl-auto=update plus custom migration runners for evolving columns/tables.")
bullet(doc, "Data model: normalized relational model with primary keys, foreign keys, unique constraints, and indexes.")
bullet(doc, "Sensitive fields: Aadhaar, PAN, bank details, phone-like fields use JPA AttributeConverter AES-GCM encryption.")

heading(doc, "4. Main Schema and Relationships")
add_table(doc, ["Table/Entity", "Purpose", "Important Relationships"], [
    ("college", "Institution tenant/master record.", "One college can have many admins/TPOs."),
    ("admin", "Institute admin account and college context.", "Belongs to college; owns courses, batches, students, teachers, fees."),
    ("course", "Course/program such as BCA/MCA.", "Many courses per admin; unique by admin and code."),
    ("batch", "Academic batch/intake.", "Belongs to course and admin."),
    ("academic_structure", "Year, semester, section structure.", "Belongs to admin, course, batch."),
    ("subject_master", "Reusable course-level subject registry.", "Belongs to course and admin."),
    ("subject", "Batch/semester subject offering.", "Belongs to course, batch, admin, optional subject master/teacher."),
    ("classroom", "Class/section/room details.", "Belongs to admin; students and timetable link to it."),
    ("student", "Student personal, academic, fee, and document state.", "Belongs to admin, batch, class."),
    ("teacher", "Teacher profile and employment data.", "Belongs to admin and optional class."),
    ("teacher_academic_mapping", "Teacher allocation to course/batch/year/subject.", "Links teacher, course, batch, subject, admin."),
    ("timetable", "Scheduled class slots.", "Links class, course, batch, subject, teacher, admin."),
    ("fees", "Fee rules and amounts.", "Belongs to admin and optional batch."),
    ("teacher_attendance_session/entry", "Attendance session and per-student entries.", "Session links teacher/class; entry links session/student."),
    ("class_assignment/submission", "Assignments and student submissions.", "Assignment links teacher/class/admin; submission links assignment/student."),
    ("exam_session/attendance/quiz_*", "Exams, quiz questions, attempts, responses.", "Exam links teacher/class/admin; attempts and responses link student/question."),
    ("study_material", "Uploaded learning materials.", "Links teacher, class, admin and file metadata."),
    ("campus_event/application", "Events and student event applications.", "Event links teacher; application links event/student."),
    ("placement_tpo/drive/application/interview", "Placement module.", "TPO creates drives; students apply; interviews map selected students."),
    ("student_document", "Uploaded admission documents.", "Links student and reviewing admin."),
    ("portal_notification", "Role/email targeted notifications.", "Stores recipient role/email, category, schedule/read status."),
    ("super_admin_audit_log", "Super-admin operation audit trail.", "Records actor, action, target, and timestamp."),
])

heading(doc, "5. Application Modules")
for module in [
    "Home and registration: public pages, college-specific registration, login routing.",
    "Super admin: college lifecycle, institute admins, communications, audit logs, reports, settings.",
    "Admin: dashboard, courses, batches, academic structure, subject registry, students, teachers, classes, timetable, fees, imports, reports.",
    "Student: dashboard, onboarding, profile, subjects, timetable, attendance, assignments, exams, fees, notifications, placements, events, study materials.",
    "Teacher: dashboard, students, subjects, attendance, marks, assignments, timetable, profile, notifications, study materials, events, exams.",
    "Placement/TPO: login, dashboard, drives, applications, students, interviews, analytics, reports, exported resumes and Excel files.",
    "Chatbot: REST context and message endpoints with role-aware conversation response objects.",
]:
    bullet(doc, module)

heading(doc, "6. Controllers and API Surface")
add_table(doc, ["Controller", "Type", "Key Routes/Responsibility"], [
    ("MainController", "MVC", "/, /register, /login, /login-student, /login-teacher, /login-admin, /superadmin, /overview, /courses, /contact."),
    ("AdminController", "MVC", "Admin dashboard, courses, batches, academic structure, imports, students, teachers, classes, timetable, fees, reports."),
    ("StudentController", "MVC", "/student-dashboard, /student-onboarding, /assignments, /profile, /fees, /notifications, /student-marks, /student-attendance."),
    ("TeacherController", "MVC", "/teacher-dashboard, /teacher-students, /teacher-attendance-submit, /teacher-marks-submit, /teacher-assignment-add."),
    ("TeacherExamController", "MVC", "/teacher-exams, /teacher-exams/create, /student-exams/{id}, quiz start/submit, quiz results."),
    ("PlacementController", "MVC", "/login-placement, /placement-dashboard, /placement-drives, /placements, /placements/apply/{driveId}, exports."),
    ("EventController", "MVC", "/events, /events/apply/{eventId}, /teacher-events, event application download."),
    ("StudyMaterialController", "MVC", "/study-materials, /teacher-study-materials/upload, view/download material."),
    ("SuperAdminController", "MVC", "/super-admin-dashboard, colleges, admins, communications, audit logs, settings, notifications, reports."),
    ("ChatbotController", "REST", "GET /api/chatbot/context and POST /api/chatbot/message returning JSON."),
])

heading(doc, "7. Security and Authentication")
doc.add_paragraph(
    "The project includes Spring Security but disables default form login, HTTP basic, and CSRF, then permits all requests at the filter-chain level. "
    "Actual page protection is implemented through SessionAccessInterceptor, which checks session attributes such as loggedInUser and userRole and redirects "
    "users to the correct login page when a protected role path is accessed without a valid session."
)
bullet(doc, "Password hashing uses BCryptPasswordEncoder with strength 12.")
bullet(doc, "Legacy password migration runner indicates older plain/legacy passwords are migrated toward protected storage.")
bullet(doc, "Sensitive database columns use AES/GCM/NoPadding encryption through SensitiveStringConverter.")
bullet(doc, "The field encryption key is configured through FIELD_ENCRYPTION_SECRET, with a development fallback value.")
bullet(doc, "Interview risk point: disabling CSRF globally is acceptable for a prototype but should be enabled for production form submissions.")

heading(doc, "8. File Uploads, Exports, and Reports")
bullet(doc, "Multipart upload limit is configured as 10 MB per file/request.")
bullet(doc, "Uploads are served from src/main/resources/static/uploads, file:uploads, and classpath static uploads.")
bullet(doc, "Student photos/documents, teacher photos, placement resumes, assignments, exam files, and study materials are stored as files with DB metadata.")
bullet(doc, "Apache POI is used for Excel import/export such as student/teacher imports, timetable export, and placement applications export.")
bullet(doc, "PDFBox and response streaming are used for PDF-oriented report/download workflows.")

heading(doc, "9. Important Interview Talking Points")
for point in [
    "Why relational DB: the domain needs joins, transactions, referential integrity, uniqueness, and reporting across related records.",
    "Why JPA repositories: reduce boilerplate CRUD and query code while keeping custom JPQL where needed.",
    "Why server-side rendering: Thymeleaf is simpler for an academic portal with forms, dashboards, and role-based pages.",
    "How multi-role access works: session role is checked by interceptor paths for ADMIN, TEACHER, STUDENT, PLACEMENT_TPO, and authenticated download routes.",
    "How data consistency is maintained: primary keys, foreign keys, unique constraints, indexes, and repository-level existence checks.",
    "How sensitive data is protected: BCrypt for passwords and AES-GCM conversion for sensitive identity/bank fields.",
    "How uploads are handled: file saved to static uploads folder; metadata such as original name, stored name, path, content type, size saved in MySQL.",
    "Current production improvements: enable CSRF, move upload storage outside source tree, externalize secrets, add stricter authorization, add Flyway/Liquibase migrations, add integration tests.",
]:
    bullet(doc, point)

heading(doc, "10. Google-Style Interview Questions and Answers")
questions = [
    ("Explain this project in 60 seconds.",
     "SCMS is a Java Spring Boot MVC application for managing college operations across super admin, admin, teacher, student, and placement roles. It uses Thymeleaf for server-rendered UI, MySQL as a relational database, Spring Data JPA/Hibernate for persistence, custom session role checks for authorization, BCrypt for password hashing, AES-GCM for sensitive field encryption, and file upload/export features for documents, resumes, assignments, timetables, and reports."),
    ("Which database did you use and why?",
     "We used MySQL, a relational database. The project has strongly connected data: colleges, admins, courses, batches, students, teachers, subjects, timetables, fees, placements, exams, and attendance. Relational modeling gives foreign keys, unique constraints, transactions, joins, and reporting capabilities, which are better fits than a document database for this domain."),
    ("What is the high-level architecture?",
     "It is a layered Spring Boot MVC application: controllers receive requests, repositories access MySQL through JPA entities, services contain domain logic for areas like academic structure, employee IDs, password protection, exams, and chatbot handling, and Thymeleaf templates render role-specific pages. Static JS/CSS supports client-side interaction."),
    ("Is it REST API based?",
     "Mostly it is an MVC web application using GET/POST controller routes that return Thymeleaf views or redirects. It also has REST-style JSON endpoints for chatbot context and messages, plus some JSON/admin utility endpoints such as options, previews, and live class data."),
    ("How are relationships modeled?",
     "JPA annotations such as @ManyToOne, @OneToMany, and @ManyToMany are used. For example, Course belongs to Admin, Batch belongs to Course/Admin, Student belongs to Admin/Batch/ClassRoom, Subject belongs to Course/Batch/Admin, PlacementDrive belongs to PlacementTpo, and PlacementInterview has a many-to-many selectedStudents relationship."),
    ("How do you prevent duplicate data?",
     "The schema uses unique constraints such as unique course code per admin, unique batch per admin/course/display name, unique subject code per admin/course/batch/semester, unique student email, and repository checks like existsByEmailIgnoreCase or existsByAdminAndEnrollmentNoIgnoreCase before saving."),
    ("How is authentication implemented?",
     "The app uses custom login routes for roles. Passwords are verified against stored hashes using BCrypt where migrated. After login, session attributes store loggedInUser and userRole. SessionAccessInterceptor maps protected URL patterns to required roles and redirects unauthenticated users."),
    ("Why is BCrypt used?",
     "BCrypt is a slow adaptive one-way password hashing algorithm. It includes salt handling and a configurable cost factor. In this project the encoder uses strength 12, making brute-force attacks more expensive than plain SHA or MD5."),
    ("How are sensitive fields protected?",
     "Fields such as Aadhaar, PAN, bank account, IFSC-like values, and some phone fields are annotated with @Convert using SensitiveStringConverter. It encrypts values before writing to DB and decrypts when reading. The encryption uses AES-GCM with a random 12-byte IV and authentication tag."),
    ("What are the biggest security improvements you would make before production?",
     "Enable CSRF protection for form submissions, replace broad permitAll with Spring Security role-based authorization, move secrets to environment/secret manager only, enforce secure cookies, validate upload content more strictly, store uploads outside source folders, add audit logging for more roles, and add rate limiting for login/chatbot endpoints."),
    ("How does file upload work?",
     "Controllers accept MultipartFile, validate/store the file under static uploads paths with generated names, and save metadata in relational tables such as StudyMaterial, StudentDocument, PlacementApplication, Assignment, AssignmentSubmission, or ExamSession. Download endpoints stream the stored file back to the user."),
    ("How would you scale this application?",
     "First separate file storage from the app server using object storage, externalize sessions to Redis or use stateless auth, add database indexes for frequent filters, introduce caching for lookup data, paginate dashboards, split large controllers into services, and deploy multiple app instances behind a load balancer."),
    ("How would you design the ER diagram?",
     "Start with College as tenant context, Admin under College, then Course, Batch, AcademicStructure, SubjectMaster, Subject, ClassRoom, Student, Teacher, and TeacherAcademicMapping. Operational modules branch from these: Timetable, Fees, Attendance, Assignment, Exam, StudyMaterial, CampusEvent, PlacementDrive/Application/Interview, Notification, and AuditLog."),
    ("What is the difference between subject_master and subject?",
     "SubjectMaster is the reusable registry for a course and semester, while Subject is the actual subject assigned/applied to a specific batch and semester. This separates curriculum definition from batch-level offering and teacher assignment."),
    ("Why use Thymeleaf instead of React?",
     "For this project, server-rendered pages reduce frontend complexity and integrate naturally with Spring MVC forms, redirects, sessions, and role-specific dashboards. React would be useful for a richer SPA, but it would require a stronger REST API boundary and client-side state management."),
    ("How are reports and exports generated?",
     "The project uses Apache POI for Excel generation/import templates and PDFBox or response streaming for PDF/download workflows. Controllers set content disposition headers and write generated files to the HTTP response."),
    ("How do repositories work in this project?",
     "Repositories extend JpaRepository and use Spring Data query derivation like findByAdminAndBatch, existsByEmailIgnoreCase, countByExamSessionAndStatusIgnoreCase, plus JPQL @Query for more specific joins and distinct filters."),
    ("What is ddl-auto=update and is it production-safe?",
     "ddl-auto=update lets Hibernate adjust schema based on entities during startup. It is convenient in development but risky in production because schema changes are implicit. Production should use versioned migration tools like Flyway or Liquibase."),
    ("How would you test this project?",
     "Use unit tests for services, repository tests with Testcontainers MySQL or @DataJpaTest, controller tests with MockMvc, integration tests for login/session role paths, upload/download tests, and security tests for unauthorized access."),
    ("What was the most complex module?",
     "Academic and role management is complex because admin-created courses, batches, academic structure, subjects, classrooms, teachers, timetables, and students must stay consistent. Placement and exam modules also add multi-entity workflows with files, status transitions, and student-specific records."),
    ("How would you explain normalization here?",
     "The schema avoids putting everything in one table. Courses, batches, subjects, students, teachers, fees, timetable, assignments, and placements are separate entities connected by keys. This reduces duplication, supports referential integrity, and makes reporting more reliable."),
    ("Where might denormalization be acceptable?",
     "Some display fields such as course name, batch name, or academic year appear on operational records to simplify historical display and filtering. The tradeoff is possible duplication, so source-of-truth fields and update rules must be clear."),
    ("What are the main indexes?",
     "Important indexes include admin, batch, class, teacher, timetable scope, fees lookup, placement drive/application links, and event/student links. These help dashboard filters and relationship joins."),
    ("How do you handle authorization for downloads?",
     "The interceptor treats file download routes as AUTHENTICATED, meaning any logged-in user can access certain downloads. For production, each download should also verify ownership or role-level permission at controller/service level."),
    ("What would you refactor first?",
     "AdminController is very large and should be split by bounded context: CourseController, StudentAdminController, TeacherAdminController, TimetableController, FeeController, ImportController, and ReportController. Business rules should move into services to improve testing and maintainability."),
]
for question, answer in questions:
    qa(doc, question, answer)

heading(doc, "11. Final Interview Summary")
doc.add_paragraph(
    "The strongest interview positioning is: SCMS is a relational, multi-role, Spring Boot MVC system with MySQL, JPA entity relationships, "
    "session-based access control, encrypted sensitive fields, file upload/export workflows, and modules that map closely to real college operations. "
    "Be ready to discuss why MySQL fits, how the schema is normalized, how each role moves through the system, what security exists today, and what you would improve for production."
)

doc.save(OUT)
print(OUT)
