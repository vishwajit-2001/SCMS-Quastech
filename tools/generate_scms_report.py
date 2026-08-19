from __future__ import annotations

import math
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:  # pragma: no cover
    Image = None
    ImageDraw = None
    ImageFont = None


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
OUTPUT = ROOT / "scms_sem_iv_project_report.docx"
ASSET_DIR = ROOT / "report_assets"
SCHEMA_PATH = SRC / "main" / "resources" / "schema.sql"
TEMPLATE_ROOT = SRC / "main" / "resources" / "templates"
JAVA_ROOT = SRC / "main" / "java" / "com" / "scms" / "scms"

TITLE = "Smart Campus Management System (SCMS)"
SUBTITLE = "Semester IV Project Report"
GROUP_MEMBERS = [
    "Vishwajit Gadale",
    "Amar Bharati",
    "Aakash Vishwakarma",
]

ROLE_DESCRIPTIONS = {
    "admin": "College administration workspace for operational management of students, teachers, courses, academic structure, fees, reports, and institutional workflows.",
    "student": "Student-facing portal for onboarding, timetable visibility, assignments, exams, placements, fee status, attendance, notifications, and profile management.",
    "teacher": "Faculty portal for classroom operations such as attendance, timetable access, assignment publishing, study material distribution, and exam management.",
    "placement": "Training and placement office workspace for company drives, student applications, interviews, exports, analytics, and placement reporting.",
    "super-admin": "Platform-wide control layer for onboarding colleges, managing college admins, reviewing audit history, and maintaining tenant isolation.",
    "home": "Public portal for institution information, contact pages, course awareness, and first-step registration access.",
    "fragments": "Reusable UI fragments that support navigation, chatbot integration, and common layout consistency across portals.",
}


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Update table in Word if page numbers do not refresh automatically."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(24)
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_after = Pt(10)

    subtitle = styles.add_style("SCMS Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor(70, 70, 70)
    subtitle.paragraph_format.space_after = Pt(12)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 16, 8),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 12, 6),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("Page ")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(90, 90, 90)
    add_page_number(footer_p)


def add_cover_page(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(TITLE)
    p = doc.add_paragraph(style="SCMS Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(SUBTITLE)
    meta = [
        "A comprehensive implementation report based on the current Spring Boot project",
        "Prepared for college submission and organizational review",
        "Academic Year: 2025-26",
        "Technology Stack: Java 17, Spring Boot 3.2.3, Thymeleaf, MySQL",
    ]
    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(12)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Group Members")
    r.bold = True
    r.font.size = Pt(13)
    for member in GROUP_MEMBERS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(member).font.size = Pt(12)
    doc.add_paragraph()
    for line in [
        "Submitted to: ________________________________",
        "Submitted by Institution/Company: ________________________________",
        "Guide/Mentor: ________________________________",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(12)
    doc.add_page_break()


def add_front_matter_page(doc: Document, heading: str, paragraphs: list[str]) -> None:
    doc.add_heading(heading, level=1)
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.add_page_break()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, True)
        set_cell_margins(hdr[idx])
        if widths:
            hdr[idx].width = widths[idx]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            set_cell_margins(cells[idx])
            if widths:
                cells[idx].width = widths[idx]
    doc.add_paragraph()


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.1) -> None:
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(width))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(10)


def paragraph_block(doc: Document, texts: list[str]) -> None:
    for text in texts:
        doc.add_paragraph(text)


def parse_schema(path: Path) -> dict[str, list[tuple[str, str]]]:
    text = path.read_text(encoding="utf-8")
    tables: dict[str, list[tuple[str, str]]] = {}
    for match in re.finditer(
        r"CREATE TABLE IF NOT EXISTS\s+([a-zA-Z0-9_]+)\s*\((.*?)\)\s*ENGINE=",
        text,
        re.S,
    ):
        table_name = match.group(1)
        body = match.group(2)
        columns: list[tuple[str, str]] = []
        for raw_line in body.splitlines():
            line = raw_line.strip().rstrip(",")
            if not line or line.upper().startswith(
                ("PRIMARY KEY", "UNIQUE KEY", "KEY ", "CONSTRAINT", ")")
            ):
                continue
            parts = line.split(None, 2)
            if len(parts) >= 2:
                columns.append((parts[0], parts[1]))
        tables[table_name] = columns
    return tables


def screen_catalog() -> dict[str, list[Path]]:
    grouped: dict[str, list[Path]] = defaultdict(list)
    for file in sorted(TEMPLATE_ROOT.rglob("*.html")):
        rel = file.relative_to(TEMPLATE_ROOT)
        grouped[rel.parts[0]].append(rel)
    return grouped


def prettify_screen_name(path: Path) -> str:
    return path.stem.replace("-", " ").replace("_", " ").title()


def screen_purpose(role: str, name: str) -> str:
    n = name.lower()
    if "dashboard" in n:
        return "acts as the operational landing page that consolidates status widgets, shortcuts, and role-specific summaries."
    if "login" in n:
        return "handles secure role entry and validates the user before routing them into the correct workspace."
    if "profile" in n:
        return "allows the role user to view or maintain their personal and institutional profile information."
    if "attendance" in n:
        return "supports attendance capture, review, or history tracking for classroom and exam operations."
    if "assignment" in n:
        return "manages assignment publication, submission, tracking, and review interactions."
    if "study" in n or "material" in n:
        return "organizes downloadable learning resources and content distribution workflows."
    if "exam" in n or "quiz" in n:
        return "supports exam lifecycle actions such as scheduling, instructions, participation, attempt review, or result visibility."
    if "timetable" in n:
        return "presents structured class scheduling information and aligns users with time-based academic activities."
    if "notification" in n or "announcement" in n:
        return "serves as a communication screen for broadcast updates, alerts, and operational notices."
    if "student" in n:
        return "manages student records, linked documents, or student-facing academic operations."
    if "teacher" in n:
        return "supports teacher-oriented administration, subject mapping, or teaching workflow coordination."
    if "course" in n or "subject" in n or "batch" in n or "academic" in n or "class" in n:
        return "maintains academic master data that drives timetable, subject allocation, class operations, and role filtering."
    if "placement" in n or "company" in n or "interview" in n:
        return "tracks placement-cycle stages from drive creation to application review and interview scheduling."
    if "event" in n:
        return "covers event publishing, participation monitoring, and activity-level academic engagement."
    if "report" in n or "analytics" in n:
        return "aggregates filtered information into management-level summaries used for review and decision making."
    if "onboarding" in n or "panel" in n:
        return "guides a staged data-entry flow so student information is captured progressively and validated in context."
    if "settings" in n:
        return "stores configurable preferences and environment-level controls for the assigned role."
    return "provides a focused workflow screen for this role within the overall campus management process."


def screen_inputs(name: str) -> str:
    n = name.lower()
    inputs = []
    if any(k in n for k in ["add", "edit", "profile", "onboarding", "student", "teacher", "course", "batch", "subject"]):
        inputs.append("validated form fields")
    if any(k in n for k in ["assignment", "study", "exam", "resume", "document", "upload", "import"]):
        inputs.append("file uploads where required")
    if any(k in n for k in ["dashboard", "report", "analytics", "fees", "attendance", "timetable", "applications"]):
        inputs.append("filters and date or status selections")
    if not inputs:
        inputs.append("role-based navigation actions")
    return ", ".join(inputs)


def screen_outputs(name: str) -> str:
    n = name.lower()
    outputs = []
    if any(k in n for k in ["dashboard", "report", "analytics"]):
        outputs.append("summary cards and reviewable metrics")
    if any(k in n for k in ["assignment", "study", "document", "resume", "import", "exam"]):
        outputs.append("downloadable or persisted file records")
    if any(k in n for k in ["students", "teachers", "courses", "subject", "batch", "notifications", "applications"]):
        outputs.append("tabular records with action controls")
    if any(k in n for k in ["onboarding", "profile", "add", "edit"]):
        outputs.append("saved profile or transaction data")
    if not outputs:
        outputs.append("screen-specific status feedback")
    return ", ".join(outputs)


def draw_box(draw, xy, text, fill, outline=(90, 90, 90), font=None):
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    width = x2 - x1
    lines = text.split("\n")
    y = y1 + 10
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        tx = x1 + (width - (bbox[2] - bbox[0])) / 2
        draw.text((tx, y), line, fill=(20, 20, 20), font=font)
        y += (bbox[3] - bbox[1]) + 6


def make_diagrams() -> dict[str, Path]:
    ASSET_DIR.mkdir(exist_ok=True)
    if Image is None:
        return {}
    font = ImageFont.load_default()
    outputs = {}

    # Architecture
    img = Image.new("RGB", (1400, 850), "white")
    d = ImageDraw.Draw(img)
    draw_box(d, (220, 70, 1180, 180), "Presentation Layer\nThymeleaf templates, CSS, JavaScript, role dashboards", (233, 242, 255), font=font)
    draw_box(d, (220, 240, 1180, 350), "Controller Layer\nMain, Admin, Student, Teacher, Placement, Event, Super Admin, Chatbot", (227, 245, 233), font=font)
    draw_box(d, (220, 410, 1180, 520), "Service Layer\nExam automation, academic structure, password protection, chatbot, identity support", (255, 243, 224), font=font)
    draw_box(d, (220, 580, 1180, 690), "Persistence Layer\nJPA repositories, entities, validation, transactional storage", (252, 228, 236), font=font)
    draw_box(d, (220, 740, 1180, 820), "Database and File Storage\nMySQL schema plus managed upload directories for documents and academic files", (240, 240, 240), font=font)
    for y1, y2 in [(180, 240), (350, 410), (520, 580), (690, 740)]:
        d.line((700, y1, 700, y2), fill=(100, 100, 100), width=3)
        d.polygon([(690, y2 - 12), (700, y2), (710, y2 - 12)], fill=(100, 100, 100))
    path = ASSET_DIR / "architecture.png"
    img.save(path)
    outputs["architecture"] = path

    # Gantt
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((60, 30), "SCMS Development Schedule (Indicative Semester Plan)", fill="black", font=font)
    phases = [
        ("Requirement Study", 1, 2),
        ("System Planning", 2, 3),
        ("Database Design", 3, 5),
        ("UI and Module Design", 4, 6),
        ("Backend Development", 5, 10),
        ("Portal Integration", 7, 11),
        ("Testing and Fixing", 10, 13),
        ("Documentation", 12, 14),
    ]
    start_x = 280
    row_y = 110
    week_w = 85
    for i in range(14):
        x = start_x + i * week_w
        d.rectangle((x, 70, x + week_w, 110), outline=(150, 150, 150))
        d.text((x + 24, 82), f"W{i+1}", fill="black", font=font)
    for idx, (label, start, end) in enumerate(phases):
        y = row_y + idx * 85
        d.text((60, y + 18), label, fill="black", font=font)
        for i in range(14):
            x = start_x + i * week_w
            d.rectangle((x, y, x + week_w, y + 50), outline=(210, 210, 210))
        x1 = start_x + (start - 1) * week_w
        x2 = start_x + end * week_w
        d.rounded_rectangle((x1 + 4, y + 6, x2 - 4, y + 44), radius=8, fill=(52, 152, 219), outline=(35, 111, 161))
    path = ASSET_DIR / "gantt.png"
    img.save(path)
    outputs["gantt"] = path

    # Use case
    img = Image.new("RGB", (1500, 950), "white")
    d = ImageDraw.Draw(img)
    d.rectangle((350, 60, 1280, 880), outline=(120, 120, 120), width=2)
    d.text((690, 75), "SCMS System Boundary", fill="black", font=font)
    actors = [("Student", (100, 180)), ("Teacher", (100, 380)), ("Admin", (100, 580)), ("Super Admin", (100, 780)), ("Placement TPO", (1360, 420))]
    for actor, (x, y) in actors:
        d.ellipse((x + 25, y, x + 75, y + 50), outline="black", width=2)
        d.line((x + 50, y + 50, x + 50, y + 120), fill="black", width=2)
        d.line((x + 20, y + 80, x + 80, y + 80), fill="black", width=2)
        d.line((x + 50, y + 120, x + 20, y + 170), fill="black", width=2)
        d.line((x + 50, y + 120, x + 80, y + 170), fill="black", width=2)
        d.text((x - 10, y + 180), actor, fill="black", font=font)
    use_cases = [
        ("Manage Registration\nand Onboarding", (520, 140, 820, 210)),
        ("View Timetable,\nFees, and Subjects", (520, 250, 820, 320)),
        ("Attendance and\nAssignments", (520, 360, 820, 430)),
        ("Exam and Quiz\nManagement", (520, 470, 820, 540)),
        ("Academic Master Data\nand Reports", (520, 580, 820, 650)),
        ("College Management,\nNotifications, Audit", (520, 690, 820, 760)),
        ("Placement Drives and\nInterview Workflow", (880, 415, 1180, 485)),
    ]
    for text, box in use_cases:
        draw_box(d, box, text, (245, 245, 245), font=font)
    line_pairs = [
        ((175, 235), (520, 175)), ((175, 435), (520, 395)), ((175, 635), (520, 615)),
        ((175, 835), (520, 725)), ((1360, 505), (1180, 450)),
    ]
    for a, b in line_pairs:
        d.line((*a, *b), fill=(100, 100, 100), width=2)
    path = ASSET_DIR / "use_case.png"
    img.save(path)
    outputs["use_case"] = path

    # Activity
    img = Image.new("RGB", (1200, 1500), "white")
    d = ImageDraw.Draw(img)
    items = [
        "Start",
        "Student registration or role login",
        "Credential and college-code validation",
        "Route to role-specific dashboard",
        "Perform academic or administrative action",
        "Persist data, files, and notifications",
        "Display confirmation and updated records",
        "End",
    ]
    y = 80
    for idx, text in enumerate(items):
        if idx in (0, len(items) - 1):
            d.ellipse((420, y, 780, y + 80), outline="black", width=2, fill=(230, 230, 230))
        else:
            draw_box(d, (330, y, 870, y + 90), text, (235, 245, 255), font=font)
        bbox = d.textbbox((0, 0), text, font=font)
        d.text((600 - (bbox[2] - bbox[0]) / 2, y + 28), text, fill="black", font=font)
        if idx < len(items) - 1:
            d.line((600, y + 80, 600, y + 130), fill=(100, 100, 100), width=3)
            d.polygon([(590, y + 118), (600, y + 130), (610, y + 118)], fill=(100, 100, 100))
        y += 150
    path = ASSET_DIR / "activity.png"
    img.save(path)
    outputs["activity"] = path

    # Class diagram
    img = Image.new("RGB", (1600, 1000), "white")
    d = ImageDraw.Draw(img)
    boxes = {
        "College": (80, 120, 330, 260),
        "Admin": (420, 120, 700, 300),
        "Student": (760, 120, 1060, 320),
        "Teacher": (1120, 120, 1440, 320),
        "ClassRoom": (420, 420, 720, 590),
        "Subject": (780, 420, 1060, 590),
        "ExamSession": (1140, 420, 1460, 590),
        "PlacementDrive": (760, 720, 1100, 880),
    }
    contents = {
        "College": "id\nname\ncode\nstatus",
        "Admin": "id\nname\nemail\ncollegeId\nrole",
        "Student": "id\nname\nemail\ncourse\nbatch\nadminId",
        "Teacher": "id\nname\nemail\nsubject\nadminId",
        "ClassRoom": "id\ncourse\nbatch\nsemester\nsection",
        "Subject": "id\ncode\nsemester\ncourseId\nteacherId",
        "ExamSession": "id\ntitle\nmode\nexamDate\nteacherId",
        "PlacementDrive": "id\ncompanyName\nroleTitle\nadminId",
    }
    for name, box in boxes.items():
        draw_box(d, box, f"{name}\n\n{contents[name]}", (250, 250, 250), font=font)
    connectors = [
        ((330, 180), (420, 180)), ((700, 180), (760, 180)), ((1060, 180), (1120, 180)),
        ((560, 300), (560, 420)), ((910, 320), (910, 420)), ((1280, 320), (1280, 420)),
        ((910, 590), (910, 720)),
    ]
    for a, b in connectors:
        d.line((*a, *b), fill=(90, 90, 90), width=3)
    path = ASSET_DIR / "class_diagram.png"
    img.save(path)
    outputs["class_diagram"] = path

    # Sequence
    img = Image.new("RGB", (1600, 950), "white")
    d = ImageDraw.Draw(img)
    lifelines = ["Student", "Portal UI", "StudentController", "Service Layer", "Repository", "Database"]
    xs = [120, 360, 620, 900, 1160, 1420]
    for x, name in zip(xs, lifelines):
        d.rectangle((x - 70, 40, x + 70, 90), outline="black", width=2, fill=(245, 245, 245))
        bbox = d.textbbox((0, 0), name, font=font)
        d.text((x - (bbox[2] - bbox[0]) / 2, 58), name, fill="black", font=font)
        d.line((x, 90, x, 850), fill=(160, 160, 160), width=2)
    messages = [
        (120, 360, 150, "Open dashboard"),
        (360, 620, 230, "GET /student/dashboard"),
        (620, 900, 310, "Load profile + timetable"),
        (900, 1160, 390, "Find student records"),
        (1160, 1420, 470, "SQL query execution"),
        (1420, 1160, 550, "Result set"),
        (1160, 900, 630, "Mapped entities"),
        (900, 620, 710, "View model"),
        (620, 360, 790, "Rendered page"),
        (360, 120, 870, "Dashboard visible"),
    ]
    for x1, x2, y, label in messages:
        d.line((x1, y, x2, y), fill=(52, 73, 94), width=3)
        d.polygon([(x2 - 10, y - 6), (x2, y), (x2 - 10, y + 6)], fill=(52, 73, 94))
        d.text(((x1 + x2) / 2 - 40, y - 22), label, fill="black", font=font)
    path = ASSET_DIR / "sequence.png"
    img.save(path)
    outputs["sequence"] = path

    # ERD
    img = Image.new("RGB", (1500, 1000), "white")
    d = ImageDraw.Draw(img)
    er_boxes = {
        "college": (80, 100, 330, 210),
        "admin": (430, 100, 680, 230),
        "course": (780, 100, 1030, 230),
        "batch": (1130, 100, 1380, 230),
        "student": (430, 350, 680, 510),
        "teacher": (780, 350, 1030, 510),
        "subject": (1130, 350, 1380, 510),
        "exam_session": (780, 680, 1080, 840),
    }
    for name, box in er_boxes.items():
        draw_box(d, box, name, (248, 248, 248), font=font)
    links = [
        ((330, 155), (430, 155)), ((680, 155), (780, 155)), ((1030, 155), (1130, 155)),
        ((555, 230), (555, 350)), ((905, 230), (905, 350)), ((1255, 230), (1255, 350)),
        ((1030, 430), (1130, 430)), ((905, 510), (905, 680)),
    ]
    for a, b in links:
        d.line((*a, *b), fill=(100, 100, 100), width=3)
    path = ASSET_DIR / "erd.png"
    img.save(path)
    outputs["erd"] = path
    return outputs


def module_data() -> list[dict[str, object]]:
    return [
        {
            "name": "Public and Registration Module",
            "actors": "Prospective students, public visitors, college administrators",
            "features": [
                "College-aware registration entry point",
                "Institution information, contact, and overview pages",
                "Controlled routing into onboarding workflows",
            ],
            "details": [
                "The public module establishes the first interaction surface for the system. It exposes registration, course awareness, overview, and contact pages while preserving the distinction between public information and protected academic operations.",
                "Because the platform supports more than one college, even the first-step registration workflow is tied to a college code or institution-aware path. This keeps the onboarding process aligned with the correct tenant from the moment the user begins interacting with the application.",
                "Operationally, this module reduces ambiguity in student intake because college-specific registration reduces the chance of a learner being attached to the wrong institution or administrative workspace.",
            ],
        },
        {
            "name": "Student Lifecycle Module",
            "actors": "Students, admin reviewers",
            "features": [
                "Multi-section onboarding with personal, academic, address, document, and payment panels",
                "Timetable, subjects, attendance, fees, placement, events, and notifications",
                "Assignment submission and exam participation",
            ],
            "details": [
                "The student module is designed around the complete life cycle of a learner after registration. It starts with guided onboarding and continues into daily academic operations such as accessing subjects, fees, attendance, study materials, and assessments.",
                "Each screen in the student portal contributes to self-service. This is important because it reduces direct dependency on office staff for routine information like classroom schedules, fee summaries, event details, and assignment status.",
                "The onboarding flow is intentionally broken into smaller connected panels so that high-volume data capture remains manageable. This improves form completion rates and makes validation more precise than a single long registration form.",
            ],
        },
        {
            "name": "Teacher Operations Module",
            "actors": "Teachers, class mentors, subject faculty",
            "features": [
                "Attendance session creation and updates",
                "Assignment publishing and study material distribution",
                "Exam scheduling, seating, quiz results, and marks",
            ],
            "details": [
                "The teacher portal converts classroom administration into a repeatable digital workflow. Rather than maintaining attendance and assignment state through disconnected files, the faculty user works within one dashboard linked to institutional master data.",
                "The module also supports subject visibility, timetable synchronization, and exam management. This means the teacher is not working in isolation; every action is tied back to class, subject, student, and academic structure records already defined by the institution.",
                "By centralizing resource uploads and academic actions, the teacher module becomes a dependable operational layer for both day-to-day teaching work and semester-level evaluation activities.",
            ],
        },
        {
            "name": "Admin Management Module",
            "actors": "College admin staff",
            "features": [
                "Student, teacher, course, batch, class, subject, and timetable management",
                "Fee setup, notifications, document review, and reporting",
                "Import center and data governance support",
            ],
            "details": [
                "The college admin module functions as the institutional control center. It handles master data, admission-linked records, teacher allocations, fee structures, and academic schedule setup.",
                "A major strength of this module is that it bridges policy-level data with operational records. For example, once a course, batch, and subject structure is defined, the rest of the system can consume those definitions for timetable, attendance, examinations, and student dashboards.",
                "This area of the application also supports reports and notifications, which means the admin interface is both transactional and supervisory. It is not merely a data entry area; it is also where institutional review happens.",
            ],
        },
        {
            "name": "Placement and Career Module",
            "actors": "Placement TPO, students, recruiting organizations",
            "features": [
                "Drive creation and application lifecycle",
                "Interview scheduling and status analytics",
                "Resume ZIP export and spreadsheet-based reporting",
            ],
            "details": [
                "Placement is handled as a dedicated operational stream instead of being treated as an afterthought. The placement workspace supports drives, applications, interviews, company coordination, and analytics relevant to campus recruitment.",
                "Students benefit because placement notices and application status become visible in a familiar portal environment, while the TPO gains workflow visibility over pipelines and exportable records.",
                "This separation into a dedicated module is particularly useful for multi-college usage because recruitment activities can vary by institution, course mix, and timeline while still remaining within the same platform architecture.",
            ],
        },
        {
            "name": "Super Admin and Platform Services Module",
            "actors": "Platform owner, system-level administrators",
            "features": [
                "College lifecycle control",
                "College-admin communication and audit log review",
                "Tenant-level oversight, security support, and chatbot integration",
            ],
            "details": [
                "The super admin module exists because the project is not limited to a single college. It provides the governance layer required to onboard, activate, pause, review, and supervise multiple institutions within one platform.",
                "This module is critical for tenant isolation because it keeps college-level administration subordinate to a controlled platform layer. That arrangement improves accountability and makes centralized maintenance feasible.",
                "Supporting services such as field encryption, schema migration runners, and contextual chatbot access extend the platform from a simple web portal into a more structured campus operations product.",
            ],
        },
    ]


def add_intro_sections(doc: Document) -> None:
    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Introduction", level=2)
    paragraph_block(
        doc,
        [
            "Smart Campus Management System (SCMS) is a centralized web-based academic administration platform built to coordinate students, teachers, college administrators, placement teams, and a platform-level super admin from a single application. The project has been implemented as a Spring Boot monolithic application using Java 17, Thymeleaf for server-rendered views, and MySQL for relational persistence.",
            "The central motivation behind the project is the fragmentation commonly seen in campus administration. In many educational institutions, classroom records, attendance logs, student onboarding forms, placement activity, and fee tracking are maintained across spreadsheets, handwritten registers, shared drives, messaging groups, and isolated software tools. That fragmentation slows decision making, increases duplication of effort, and makes role accountability difficult.",
            "SCMS addresses this problem by creating a role-based digital platform in which every major academic workflow has a defined owner, storage model, and user interface. The same application can serve multiple colleges while still preserving institution-specific data isolation through college-aware routing and administration.",
            "From an implementation perspective, the project demonstrates how a structured campus domain can be translated into modular controllers, service classes, repositories, entities, templates, file storage, and security configuration. From a business perspective, it shows how academic and administrative operations can move from ad hoc manual execution to trackable digital processes.",
        ],
    )

    doc.add_heading("1.1.1 Problem Definition", level=3)
    paragraph_block(
        doc,
        [
            "Educational institutions frequently face operational inefficiencies because information is scattered across departments. Student admission details may be recorded in one place, attendance in another, fee records elsewhere, and placement workflows in disconnected spreadsheets or email chains. When different teams work in isolation, the institution loses real-time visibility into student progress and administrative status.",
            "The absence of a common digital workflow also introduces quality issues. Duplicate data entry, missing files, delayed approvals, inconsistent timetable communication, and weak reporting practices can directly affect students, teachers, and office staff. Even when institutions adopt software, those tools are often siloed and do not cover the full lifecycle from registration to placement.",
            "The project therefore defines the core problem as the need for one integrated, role-aware campus platform that can connect admissions, academics, teaching operations, fee management, events, notifications, examinations, and placement processes while preserving accountability and reducing manual overhead.",
        ],
    )
    add_bullets(
        doc,
        [
            "Manual and repeated data entry across departments wastes staff time.",
            "Students depend too heavily on offline communication for routine information.",
            "Teachers lack one unified place to manage attendance, resources, and assessments.",
            "Admin teams struggle to maintain synchronized master data and institution-wide reports.",
            "Placement activities are difficult to monitor when applications and interviews are tracked in isolated documents.",
            "Without platform-level oversight, multi-college deployment becomes inconsistent and difficult to audit.",
        ],
    )

    doc.add_heading("1.1.2 Objectives of Project", level=3)
    paragraph_block(
        doc,
        [
            "The objective of SCMS is not simply to computerize a few forms. The broader goal is to build a dependable operational system that supports day-to-day campus activity while remaining extensible enough for institutional growth. The project therefore combines domain modeling, access control, file handling, reporting support, and guided user interfaces into a single cohesive platform.",
            "Each objective below has been chosen to reflect both technical and institutional value. Some objectives focus on usability and operational efficiency, while others focus on data integrity, security, and maintainability.",
        ],
    )
    add_numbered(
        doc,
        [
            "To provide a unified platform for student, teacher, admin, placement, and super admin operations.",
            "To digitize student onboarding through a structured multi-panel registration completion workflow.",
            "To manage academic structure in a reusable and institution-linked manner using courses, batches, classes, and subjects.",
            "To support classroom operations such as attendance, timetable visibility, assignments, and study materials.",
            "To manage examinations and PDF-assisted quiz generation in a role-aware digital workflow.",
            "To maintain fee, notification, event, and profile data within student-facing self-service dashboards.",
            "To support placement drives, interviews, resume handling, and reporting from a dedicated TPO portal.",
            "To introduce platform-level governance for multi-college onboarding and audit support.",
            "To protect sensitive data through password hashing, access control, and field encryption support.",
            "To ensure the solution remains modular and maintainable through controllers, services, repositories, and entities.",
        ],
    )

    doc.add_heading("1.1.3 Scope of Project", level=3)
    paragraph_block(
        doc,
        [
            "The current scope of SCMS covers the operational boundary that is both useful and realistically implemented in the codebase. The project is intentionally broad enough to serve as a complete campus workflow demonstration, but it also avoids overstating capabilities that are not yet implemented.",
            "The system focuses on browser-based usage through a role-driven web application. It does not presently attempt to solve every possible institutional integration problem, nor does it present itself as a fully distributed or AI-predictive platform.",
        ],
    )
    add_simple_table(
        doc,
        ["Scope Area", "Current Coverage"],
        [
            ["In Scope", "College-aware registration, onboarding, master data, attendance, assignments, study materials, exams, fees, events, notifications, placement workflows, audit support, and file handling."],
            ["Technical Scope", "Spring Boot web application, Thymeleaf templates, MySQL persistence, session-aware role routing, encryption support, and static upload storage."],
            ["Out of Scope", "Native mobile apps, biometric hardware integration, external university APIs, predictive machine learning models, and production-grade microservice decomposition."],
        ],
    )

    doc.add_heading("1.2 Technical Details", level=2)
    paragraph_block(
        doc,
        [
            "SCMS has been implemented as a modern Java web application using a layered architecture. The technology stack has been selected to keep development productive while still supporting real-world domain complexity such as authentication, relational persistence, multipart upload handling, and structured UI rendering.",
            "The backend framework is Spring Boot 3.2.3 on Java 17. Persistence is handled by Spring Data JPA backed by MySQL. Security is implemented through BCrypt password hashing and session-based route protection. Thymeleaf templates and static CSS/JavaScript deliver the user interface. Apache POI and PDFBox support file-driven operations such as spreadsheet export and PDF-based content extraction.",
        ],
    )

    doc.add_heading("1.2.1 Overview of the Front End", level=3)
    paragraph_block(
        doc,
        [
            "The front end is implemented through server-rendered Thymeleaf templates backed by static CSS and JavaScript assets. This approach keeps the rendering model simple while still allowing differentiated dashboards and workflows for each role.",
            "A notable design characteristic is the separation of templates by role. Student, teacher, admin, placement, and super admin users each receive their own template directories, which keeps navigation clear and reduces the risk of mixing role responsibilities at the UI layer.",
            "The onboarding flow is especially important from a front-end perspective. Rather than forcing the student to fill an unstructured form, the application breaks the process into personal, address, education, document, payment, and confirmation panels. This improves usability and supports more targeted validation.",
            "Template fragments, sidebars, and shared components allow repeated interface patterns to remain consistent. This reduces maintenance effort because navigation, badges, and layout support are not duplicated across every HTML screen.",
        ],
    )
    add_simple_table(
        doc,
        ["Front-End Aspect", "Implementation Note"],
        [
            ["Rendering model", "Server-side HTML generation with Thymeleaf"],
            ["Styling", "Role-specific CSS plus reusable common styles"],
            ["Interaction", "JavaScript helpers for onboarding, dashboard behavior, and client-side validation"],
            ["Layout principle", "Separate template folders for each portal role"],
            ["Usability emphasis", "Dashboard-led navigation and staged workflows for complex forms"],
        ],
    )

    doc.add_heading("1.2.2 Overview of the Back End", level=3)
    paragraph_block(
        doc,
        [
            "The backend follows a layered Spring Boot structure. Controllers handle HTTP requests and route the user to the appropriate template or business operation. Services encapsulate reusable logic such as exam automation, identity support, academic mapping, password protection, and chatbot response generation.",
            "Repositories provide persistence access through Spring Data JPA, while entity classes model the institutional domain. This combination keeps business rules separate from request mapping and supports a maintainable structure as the project grows.",
            "The application also includes operational support classes such as schema migration runners, encryption helpers, security configuration, and session interceptors. These components are important because they show that the system is intended for evolving real data rather than being a purely static academic demonstration.",
            "File storage is handled through managed static directories so that uploaded student documents, profile images, assignments, resumes, and exam files remain available to the portal experience. This makes the application both data-centric and document-centric.",
        ],
    )


def add_system_study(doc: Document, diagrams: dict[str, Path]) -> None:
    doc.add_heading("2. System Study and Planning", level=1)
    doc.add_heading("2.1 System Study", level=2)
    paragraph_block(
        doc,
        [
            "System study was carried out to understand how academic administration is executed in practice and where the most significant friction points occur. The study considered student intake, class organization, timetable visibility, attendance, internal communication, placement workflow, and multi-college governance requirements.",
            "The result of the study was that institutions often do not suffer from a lack of data; instead, they suffer from a lack of consolidated workflow. The same information exists in multiple places, but there is no common operational path through which all actors can work consistently.",
        ],
    )
    doc.add_heading("2.1.1 Existing System", level=3)
    paragraph_block(
        doc,
        [
            "In a conventional campus administration environment, many activities are managed manually or through disconnected tools. Student admission details are often collected using forms and spreadsheets. Attendance may remain inside classroom registers or department-owned files. Fee status is maintained in a separate office process. Placement information moves through messaging groups, email, or standalone spreadsheets.",
            "These isolated practices produce partial visibility at best. A teacher may not easily access the same up-to-date student profile information that the office holds. A placement team may not have immediate access to student eligibility information. A student may depend on direct office visits to understand fee, exam, or timetable status.",
        ],
    )
    doc.add_heading("2.1.2 Disadvantages of Existing System", level=3)
    add_bullets(
        doc,
        [
            "Repeated manual entry increases errors and consumes office time.",
            "There is no unified student record that travels across academic and administrative workflows.",
            "Communication delays arise because notices are distributed through ad hoc channels.",
            "Reporting quality remains weak because data has to be consolidated manually.",
            "Security and traceability are limited when records are shared informally.",
            "Growth to multi-college administration becomes difficult without a common platform.",
        ],
    )
    paragraph_block(
        doc,
        [
            "From a planning standpoint, these disadvantages justified an integrated design rather than a partial one. If only one or two modules had been digitized, the institution would still face process fragmentation. Therefore, the proposed solution was framed as a unified campus workflow system.",
        ],
    )
    doc.add_heading("2.1.3 Proposed System", level=3)
    paragraph_block(
        doc,
        [
            "The proposed system is a role-based web platform in which students, teachers, college administrators, placement officers, and a platform super admin each interact with the same institutional data model through tailored interfaces. This preserves both consistency and separation of responsibility.",
            "The proposed design combines academic data, user profile data, operational records, uploaded documents, and supervisory records into a single structured system. Instead of routing users through generic dashboards, the system assigns them role-specific portals backed by shared business entities and scoped access rules.",
            "The proposal also includes a governance layer for multi-college oversight. This is strategically important because it transforms the project from a single-campus tool into a platform-oriented system that can scale organizationally.",
        ],
    )
    add_simple_table(
        doc,
        ["Dimension", "Existing Practice", "SCMS Proposed Approach"],
        [
            ["Student data", "Separate forms and files", "Unified digital onboarding and profile"],
            ["Attendance", "Manual or isolated records", "Timetable-linked attendance workflow"],
            ["Assignments and materials", "Messaging or informal sharing", "Teacher upload and student access portal"],
            ["Placement", "Spreadsheets and email coordination", "Dedicated TPO dashboard with exports and tracking"],
            ["Governance", "College-level silos", "Super admin layer with audit support"],
        ],
    )

    doc.add_heading("2.2 System Planning and Schedule", level=2)
    paragraph_block(
        doc,
        [
            "System planning for SCMS was based on phased delivery. Rather than attempting to complete all modules at once, the project is better understood as a sequence of requirement study, data modeling, portal implementation, integration, testing, and documentation stages.",
            "The planning assumption was that core architectural decisions such as role separation, normalized data design, and session-based security had to be established early. Once those foundations were in place, iterative module development became more manageable.",
        ],
    )
    doc.add_heading("2.2.1 S/W Development Model", level=3)
    paragraph_block(
        doc,
        [
            "The most suitable development approach for this project is an iterative and incremental model with strong alignment to practical agile behavior. The project domain is broad, and many modules are interconnected. Because of that, it is useful to stabilize major architecture early while implementing functional slices over time.",
            "An iterative model is appropriate because student, teacher, admin, placement, and super admin workflows can be expanded progressively. Early iterations establish login, role routing, and master data. Later iterations introduce attendance, study materials, exams, placements, notifications, and automation-oriented support services.",
            "This model also matches the reality of academic projects and internship-based development, where requirements become clearer as the developer works through real screens, entities, and user interactions.",
        ],
    )
    add_bullets(
        doc,
        [
            "Requirement gathering and problem definition",
            "Data modeling and architectural decomposition",
            "Portal-specific module implementation",
            "Integration of files, reports, and notifications",
            "Testing, defect fixing, and documentation refinement",
        ],
    )
    doc.add_heading("2.2.2 GANTT Chart", level=3)
    paragraph_block(
        doc,
        [
            "The following Gantt-style representation shows an indicative semester delivery schedule. It reflects the logical ordering of discovery, design, implementation, integration, testing, and reporting activities for the project.",
        ],
    )
    add_figure(doc, diagrams.get("gantt", Path()), "Figure 2.1: Indicative semester development schedule for SCMS")
    add_simple_table(
        doc,
        ["Milestone", "Outcome"],
        [
            ["Requirement completion", "Functional scope clarified for role-based campus workflows"],
            ["Architecture and schema baseline", "Core entities and package structure stabilized"],
            ["Portal development", "Student, teacher, admin, placement, and super admin modules implemented"],
            ["Integration pass", "Uploads, reports, notifications, and exam workflows linked"],
            ["Testing and documentation", "System validation and formal project reporting completed"],
        ],
    )


def add_system_design(doc: Document, diagrams: dict[str, Path], tables: dict[str, list[tuple[str, str]]], screens: dict[str, list[Path]]) -> None:
    doc.add_heading("3. System Design", level=1)
    doc.add_heading("3.1 Software Requirement Specification (SRS)", level=2)
    doc.add_heading("3.1.1 Introduction of SRS", level=3)
    paragraph_block(
        doc,
        [
            "The Software Requirement Specification defines the functional and technical expectations that the implemented SCMS platform must satisfy. Because the application spans multiple user roles, the SRS acts as the common reference point that links business objectives to interface behavior, backend logic, and data design.",
            "In this project, the SRS is especially important because the system is not a single-purpose portal. It combines admissions-oriented onboarding, academic operations, assessments, communications, placement management, and administrative governance. A clear specification is therefore necessary to preserve coherence across modules.",
            "The implemented project aligns with the following SRS perspective: the product is a browser-based enterprise campus management application, delivered as a single Spring Boot system, with college-aware routing and role-specific views layered over a shared institutional data model.",
        ],
    )

    doc.add_heading("3.1.2 Technology Requirements", level=3)
    paragraph_block(
        doc,
        [
            "Technology requirements were selected to ensure that the system could support relational persistence, secure authentication, modular Java development, file handling, and predictable view rendering. The chosen stack also reflects technologies that are practical for academic projects while still being relevant for real-world deployment paths.",
        ],
    )
    doc.add_heading("3.1.2.1 Hardware to be Used", level=3)
    add_simple_table(
        doc,
        ["Hardware Element", "Recommended Specification", "Reason"],
        [
            ["Development laptop/desktop", "Intel i5 / Ryzen 5 or above, 8 GB RAM minimum", "Supports Java IDE, local MySQL, browser testing, and build execution"],
            ["Storage", "At least 256 GB SSD", "Required for source code, database files, and uploaded assets"],
            ["Network access", "Stable internet/LAN during setup and collaboration", "Useful for dependency retrieval, documentation, and deployment testing"],
            ["Backup media", "External/cloud backup recommended", "Protects code, documents, and exported reports"],
        ],
    )

    doc.add_heading("3.1.2.2 Software/Tools to be Used", level=3)
    add_simple_table(
        doc,
        ["Tool", "Usage in Project"],
        [
            ["Java 17", "Primary language runtime"],
            ["Spring Boot 3.2.3", "Web application framework"],
            ["Spring MVC + Thymeleaf", "Template-based portal rendering"],
            ["Spring Data JPA / Hibernate", "Persistence mapping and repository abstraction"],
            ["MySQL", "Relational database storage"],
            ["Maven", "Build and dependency management"],
            ["Apache POI", "Spreadsheet import/export workflows"],
            ["Apache PDFBox", "PDF parsing for quiz-support functionality"],
            ["Visual Studio Code / IntelliJ IDEA", "Development environment"],
            ["Git", "Version control"],
            ["LibreOffice / Word", "Documentation rendering and review"],
        ],
    )

    doc.add_heading("3.2 Detailed Life Cycle of the Project", level=2)
    paragraph_block(
        doc,
        [
            "The life cycle of SCMS can be described as a progression from problem discovery to structured implementation and validation. The project began with identifying fragmented academic administration problems, then moved through architecture design, entity modeling, portal creation, integration, testing, and documentation.",
            "Because the system covers multiple domain areas, each life-cycle step required balancing user experience concerns with backend data integrity. For example, onboarding design had to remain simple for students while still collecting enough information for institutional administration.",
        ],
    )

    doc.add_heading("3.2.1 Modules", level=3)
    paragraph_block(
        doc,
        [
            "The system is divided into modular role-oriented and function-oriented segments. This modularization keeps responsibilities clear and makes future maintenance more realistic. The modules below summarize the major functional areas implemented in the current codebase.",
        ],
    )
    for module in module_data():
        doc.add_heading(module["name"], level=3)
        paragraph_block(doc, module["details"])
        doc.add_paragraph(f"Primary actors: {module['actors']}")
        add_bullets(doc, list(module["features"]))

    doc.add_heading("3.2.2 Object Oriented Analysis & Design Diagrams", level=3)
    paragraph_block(
        doc,
        [
            "Object oriented analysis and design diagrams are used to explain the system from multiple viewpoints. Each diagram emphasizes a different form of understanding: actors and interactions, activity flow, static structure, request sequencing, and data relationships.",
        ],
    )

    doc.add_heading("3.2.2.1 Use Case Diagram", level=3)
    add_figure(doc, diagrams.get("use_case", Path()), "Figure 3.1: Use case view of the major SCMS roles and services")
    paragraph_block(
        doc,
        [
            "The use case diagram shows that the platform is not centered on a single user class. Instead, each role works with the system through a different set of responsibilities. Students consume academic and administrative information, teachers manage instructional workflows, admins govern institution-level data, placement officers manage recruitment processes, and the super admin oversees tenant-level control.",
            "This role separation is directly reflected in the implementation through dedicated controllers, templates, and dashboard structures.",
        ],
    )

    doc.add_heading("3.2.2.2 Activity Diagram", level=3)
    add_figure(doc, diagrams.get("activity", Path()), "Figure 3.2: Generic activity flow from login to completed transaction")
    paragraph_block(
        doc,
        [
            "The activity view captures the high-level path from authentication to action completion. Although different modules execute different business logic, they share a common operational pattern: validate user, resolve scope, perform the business operation, persist the result, and render feedback.",
            "This common flow improves maintainability because controllers and services can be designed around repeatable expectations.",
        ],
    )

    doc.add_heading("3.2.2.3 Class Diagram", level=3)
    add_figure(doc, diagrams.get("class_diagram", Path()), "Figure 3.3: Simplified class diagram for key domain entities")
    paragraph_block(
        doc,
        [
            "The class diagram highlights representative relationships across the institutional domain. Colleges connect to admins; admins supervise courses, students, and teachers; classrooms and subjects organize academic structure; and specialized operational entities such as exams and placements extend that base model.",
            "The actual implementation contains more entities than shown in the simplified figure, but the diagram accurately reflects how the design is centered on shared identifiers and institution-linked records.",
        ],
    )

    doc.add_heading("3.2.2.4 Sequence Diagram", level=3)
    add_figure(doc, diagrams.get("sequence", Path()), "Figure 3.4: Example request sequence for student dashboard access")
    paragraph_block(
        doc,
        [
            "The sequence view demonstrates how a request travels through the portal UI, controller, service layer, repository abstraction, and database before the rendered page returns to the user. This helps explain why the layered architecture was chosen.",
            "It also shows that the application remains understandable even though it covers many modules. Each request still follows a defined, inspectable path through the stack.",
        ],
    )

    doc.add_heading("3.2.2.5 Flowchart/DFD/ER Diagram", level=3)
    add_figure(doc, diagrams.get("erd", Path()), "Figure 3.5: Simplified ER-style data relationship overview")
    paragraph_block(
        doc,
        [
            "From a data-flow perspective, SCMS collects inputs from users through role-specific forms and then routes the resulting records into normalized relational tables. The ER-style diagram emphasizes that major records are not independent. Student operations depend on academic structure and admin scope, while exam and placement records depend on earlier master data.",
            "This relationship-aware approach reduces inconsistency and makes reporting more reliable than independent flat-file tracking.",
        ],
    )

    doc.add_heading("3.2.3 Database", level=3)
    add_figure(doc, diagrams.get("architecture", Path()), "Figure 3.6: Layered technical architecture supporting the database-backed campus workflow")
    paragraph_block(
        doc,
        [
            "The SCMS database is implemented in MySQL and follows a normalized design. Its responsibility is not limited to storing user login data; it serves as the structural backbone for onboarding, academic organization, attendance, assignments, fees, exams, events, placement processes, notifications, and audit support.",
            "A strong database design is important in a project like this because many workflows overlap. Students belong to courses and batches, classes connect to timetables and teachers, subjects influence assessments, and placement workflows rely on student and institutional context.",
            "The database therefore acts as the source of truth across the campus lifecycle rather than a passive storage component.",
        ],
    )
    add_bullets(
        doc,
        [
            "Normalized entities reduce duplication and preserve relationship clarity.",
            "Tenant awareness is preserved through college-linked administrative scope.",
            "Operational data is separated from master data for maintainability.",
            "Sensitive information can be protected through encryption support.",
            "The schema supports long-lived academic processes instead of isolated transactions.",
        ],
    )

    doc.add_heading("3.2.3.1 Database Table", level=3)
    paragraph_block(
        doc,
        [
            "The following data dictionary summarizes the primary database tables currently represented in the schema and model layer. These tables demonstrate how the system has been decomposed into core institutional records and process-specific records.",
        ],
    )
    for table_name, columns in tables.items():
        doc.add_heading(table_name, level=3)
        purpose = {
            "college": "Stores institution-level identity and status used for multi-college administration.",
            "admin": "Stores college admin accounts and their link to a managed institution.",
            "student": "Stores learner profile, academic, contact, and role-linked information.",
            "teacher": "Stores faculty identity, employment, and teaching-linked profile information.",
            "placement_tpo": "Stores placement office users and college-linked placement authority.",
        }.get(table_name, f"Stores operational or master data related to the {table_name.replace('_', ' ')} workflow.")
        doc.add_paragraph(purpose)
        rows = []
        for col, dtype in columns[:10]:
            rows.append([col, dtype])
        add_simple_table(doc, ["Column", "Type"], rows)
        if len(columns) > 10:
            doc.add_paragraph(
                f"Additional fields are present in the implemented schema. The table currently contains {len(columns)} columns, which indicates the level of detail required for this module."
            )
        doc.add_paragraph(
            f"Design note: `{table_name}` contributes to the relational backbone of SCMS and participates in controlled persistence rather than ad hoc file-based storage."
        )

    doc.add_heading("3.2.4 I/O Screen Layout", level=3)
    paragraph_block(
        doc,
        [
            "User interface design in SCMS is organized by role and workflow. The screen catalog below summarizes the current template layout implemented in the project. Rather than listing only a few example pages, this section documents the breadth of the actual portal surface so that the report accurately reflects the submitted system.",
        ],
    )
    for role, files in screens.items():
        role_desc = ROLE_DESCRIPTIONS.get(role, "Supporting template group for the platform.")
        doc.add_heading(f"{role.title()} Screen Group", level=3)
        doc.add_paragraph(role_desc)
        for rel in files:
            name = prettify_screen_name(rel)
            doc.add_heading(name, level=3)
            doc.add_paragraph(
                f"This screen belongs to the `{role}` interface family and {screen_purpose(role, rel.stem)}"
            )
            add_simple_table(
                doc,
                ["Aspect", "Details"],
                [
                    ["Template path", str(rel).replace("\\", "/")],
                    ["Primary inputs", screen_inputs(rel.stem)],
                    ["Primary outputs", screen_outputs(rel.stem)],
                    ["Role context", role],
                ],
            )
            doc.add_paragraph(
                "From a usability perspective, this screen contributes to the platform by reducing context switching and keeping related information close to the actor who needs it. The role-specific template structure also supports maintainability because UI behavior can evolve without flattening all screens into one generic layout."
            )


def add_coding_section(doc: Document, controllers: list[str], models: list[str]) -> None:
    doc.add_heading("4. Coding", level=1)
    paragraph_block(
        doc,
        [
            "The coding phase transformed the design into a structured Spring Boot application. Rather than treating coding as simple screen development, the project implemented a layered application with domain entities, repositories, services, security support, migration runners, and role-based template rendering.",
            "The implementation pattern chosen for SCMS emphasizes clarity. Controllers are responsible for routing and request handling, services encapsulate reusable business logic, repositories manage persistence access, and entities define the institutional domain model.",
        ],
    )
    add_simple_table(
        doc,
        ["Package Area", "Purpose"],
        [
            ["controller", "Handles request mapping, role routing, and view preparation"],
            ["service", "Contains reusable operational logic"],
            ["repository", "Provides persistence access through Spring Data JPA"],
            ["model", "Defines the relational domain model"],
            ["config/security", "Supports route protection, interceptors, encryption, and migration runners"],
            ["templates/static", "Delivers role-based UI screens and client assets"],
        ],
    )

    doc.add_heading("Controller Layer Implementation", level=2)
    doc.add_paragraph(
        "The controller set demonstrates that the project has been divided according to domain and role boundaries. The following controllers are present in the implementation:"
    )
    add_bullets(doc, controllers)
    paragraph_block(
        doc,
        [
            "This controller decomposition is significant because it limits the size of any one entry point class and keeps portal logic understandable. MainController handles registration and general navigation. StudentController and TeacherController focus on academic workflows. AdminController supervises institutional management. PlacementController handles recruitment operations, while SuperAdminController governs multi-college administration.",
            "TeacherExamController and EventController further show that the codebase was allowed to evolve by feature area instead of forcing unrelated responsibilities into one large class. This improves readability and testing feasibility.",
        ],
    )

    doc.add_heading("Model and Persistence Layer", level=2)
    doc.add_paragraph(
        "The current model package includes a broad set of entities representing users, academics, assessments, placement workflows, communications, and institutional control records."
    )
    add_bullets(doc, models)
    paragraph_block(
        doc,
        [
            "This entity set indicates a realistic domain scope. The project is not limited to a few academic tables; it includes supporting records such as notifications, audit logs, document storage references, attendance sessions, quiz responses, and placement interviews.",
            "Because these entities are backed by repositories, the system can execute structured queries and persist relationships in a maintainable way. That approach is far stronger than directly writing SQL inside controllers for every operation.",
        ],
    )

    doc.add_heading("Security and Data Protection", level=2)
    paragraph_block(
        doc,
        [
            "Security in SCMS is handled through multiple layers. Passwords are protected with BCrypt hashing, which prevents plain-text credential storage. Session-backed route protection ensures that a user cannot simply navigate into another role's protected pages without valid context.",
            "The project also includes field encryption support for sensitive values such as Aadhaar, bank account, and related identity or financial fields. This is an important design decision because educational systems often store regulated personal information.",
            "Path normalization and controlled file handling are also relevant from a security perspective. Since the application supports document uploads and downloads, safe file access rules are necessary to reduce traversal-related risk.",
        ],
    )

    doc.add_heading("File Handling and Utility Services", level=2)
    paragraph_block(
        doc,
        [
            "A practical campus system must manage documents as well as database records. SCMS therefore stores academic files, resumes, assignments, exam materials, student documents, and profile assets through structured upload directories.",
            "Utilities such as Apache POI and PDFBox extend the project beyond CRUD operations. POI is used for spreadsheet-style export and import support, while PDFBox enables the automation-oriented exam flow in which PDF study material can feed quiz preparation logic.",
            "These utility integrations help the platform support actual office work rather than only academic demonstrations of login and database access.",
        ],
    )

    doc.add_heading("Maintainability and Evolution", level=2)
    paragraph_block(
        doc,
        [
            "The project includes migration runners and environment-driven configuration, which indicates that schema and security evolution were considered during implementation. That is an important software engineering quality signal because campus systems rarely remain static after first deployment.",
            "Template separation, modular Java packaging, repository use, and service extraction collectively make the codebase easier to maintain. They also reduce the cost of future enhancements such as API expansion, deployment hardening, or integration with new campus subsystems.",
        ],
    )


def add_testing_section(doc: Document) -> None:
    doc.add_heading("5. Testing", level=1)
    paragraph_block(
        doc,
        [
            "Testing for SCMS must verify not only whether the application runs, but whether each role can perform the correct actions with correct data visibility and protected access. Because the platform covers multiple workflows, a structured testing strategy is necessary.",
            "The testing approach for this project combines methodology-driven validation with scenario-oriented checks. The main objective is to ensure that user access, academic transactions, document handling, reporting, and workflow transitions behave correctly.",
        ],
    )
    doc.add_heading("5.1 Methodologies Used for Testing", level=2)
    paragraph_block(
        doc,
        [
            "Testing methodology for SCMS aligns with layered application verification. Unit-like logic checks are relevant for service behavior, integration checks are relevant for controller-to-repository flow, and system testing is necessary for end-to-end portal behavior.",
            "Because the project is a web-based system with multiple actors, black-box testing and role-based scenario testing are particularly important. The tester must verify not only correct output, but also correct role restrictions and user experience continuity.",
        ],
    )
    add_bullets(
        doc,
        [
            "Requirement-based testing against the defined functional scope",
            "Role-wise scenario testing for student, teacher, admin, placement, and super admin users",
            "Form validation and error-path testing for onboarding and master-data screens",
            "Integration testing for controller, service, repository, and database flow",
            "Regression testing after feature additions and bug fixes",
        ],
    )

    doc.add_heading("5.2 Types of Testing (Whichever Used)", level=2)
    paragraph_block(
        doc,
        [
            "The system is suitable for multiple complementary forms of testing. Each type verifies a different quality dimension and together they provide stronger confidence in the product.",
        ],
    )
    add_simple_table(
        doc,
        ["Testing Type", "Purpose in SCMS"],
        [
            ["Unit Testing", "Validates isolated logic components where feasible, especially helper and service behavior."],
            ["Integration Testing", "Confirms data flow across controllers, services, repositories, and persistence."],
            ["System Testing", "Checks complete user journeys across the web application."],
            ["User Acceptance Testing", "Ensures that institutional roles can complete expected tasks practically."],
            ["Security Testing", "Reviews access control, password handling, and file access safety."],
            ["Regression Testing", "Confirms that new changes do not break existing workflows."],
            ["Usability Testing", "Assesses whether staged workflows and dashboards remain understandable to users."],
        ],
    )
    doc.add_paragraph("Representative test cases used for report documentation are listed below.")
    add_simple_table(
        doc,
        ["Test ID", "Scenario", "Expected Result"],
        [
            ["TC-01", "Student registers with valid institution path", "Account is created and onboarding access becomes available."],
            ["TC-02", "Student submits incomplete onboarding section", "Validation error is shown and no invalid data is committed."],
            ["TC-03", "Teacher opens timetable dashboard", "Assigned timetable and subject context are visible."],
            ["TC-04", "Teacher uploads assignment file", "Assignment record and file metadata persist correctly."],
            ["TC-05", "Admin creates course and batch", "Records become visible for later academic structure linkage."],
            ["TC-06", "Admin opens fee management", "Configured fee records load and remain scoped to the institution."],
            ["TC-07", "Placement TPO creates drive", "Drive persists and becomes available for application tracking."],
            ["TC-08", "Super admin changes college status", "Status is updated and institution state reflects the change."],
            ["TC-09", "Unauthorized user enters protected route", "Access is denied or redirected safely."],
            ["TC-10", "Student opens exam result page", "Only relevant exam result information is displayed."],
            ["TC-11", "Placement export generates output", "Downloadable spreadsheet or ZIP output is created correctly."],
            ["TC-12", "Notification is sent to target role", "Target user sees the notification in the correct portal area."],
        ],
    )
    paragraph_block(
        doc,
        [
            "Beyond these representative cases, the project also benefits from non-functional verification such as checking upload behavior with moderate file sizes, verifying dashboard responsiveness under realistic data volumes, and ensuring that static asset paths remain valid in development-time execution.",
            "The testing chapter is intentionally broad because a campus platform cannot be declared complete through a small number of happy-path checks. Multi-role consistency and data correctness matter as much as raw feature availability.",
        ],
    )


def add_closeout_sections(doc: Document) -> None:
    doc.add_heading("6. Conclusion", level=1)
    paragraph_block(
        doc,
        [
            "SCMS demonstrates how a broad institutional problem can be translated into a single cohesive software system. The project unifies student onboarding, teacher operations, administrative control, placement workflow, notifications, file handling, and super admin governance inside one structured platform.",
            "Technically, the project shows effective use of Spring Boot, Thymeleaf, JPA, MySQL, and supporting libraries to solve a realistic campus domain. Functionally, it reduces manual fragmentation and improves access to academic and administrative information.",
            "The system is therefore meaningful both as a semester project and as a foundation for further industrial-grade refinement.",
        ],
    )

    doc.add_heading("7. Limitations", level=1)
    add_bullets(
        doc,
        [
            "The current implementation is monolithic and not yet decomposed for distributed deployment.",
            "A native mobile application experience is not included in the present scope.",
            "The chatbot and automation features are assistive but not based on a full predictive ML pipeline.",
            "External institutional integrations such as ERP, biometric devices, or university APIs are not yet implemented.",
            "Production hardening for large-scale deployment would require additional infrastructure, monitoring, and operational controls.",
        ],
    )

    doc.add_heading("8. Future Enhancements", level=1)
    paragraph_block(
        doc,
        [
            "The future scope of SCMS is substantial because the present architecture already provides a usable operational foundation. Expansion can therefore focus on scale, automation, integration, and user convenience rather than rebuilding the domain model from scratch.",
        ],
    )
    add_numbered(
        doc,
        [
            "Add native mobile applications or responsive PWA enhancements for student and teacher access.",
            "Introduce advanced analytics for attendance risk, fee trends, placement readiness, and academic performance.",
            "Integrate with external university services, payment gateways, biometric attendance devices, or document-verification APIs.",
            "Expand the chatbot into a stronger AI assistant with secure knowledge retrieval and guided campus actions.",
            "Separate high-load modules into microservice-ready components where required.",
            "Add richer audit, backup, monitoring, and deployment automation for production environments.",
        ],
    )

    doc.add_heading("9. References", level=1)
    add_numbered(
        doc,
        [
            "Spring Boot Official Documentation.",
            "Spring Security Reference Documentation.",
            "Spring Data JPA Reference Guide.",
            "Thymeleaf Official Documentation.",
            "MySQL 8.0 Reference Manual.",
            "Apache POI Documentation.",
            "Apache PDFBox Documentation.",
            "Project source code and schema files available in the submitted SCMS workspace.",
        ],
    )


def build_report() -> None:
    doc = Document()
    configure_document(doc)

    tables = parse_schema(SCHEMA_PATH)
    screens = screen_catalog()
    diagrams = make_diagrams()
    controllers = [p.name for p in sorted((JAVA_ROOT / "controller").glob("*.java"))]
    models = [p.name for p in sorted((JAVA_ROOT / "model").glob("*.java"))]

    add_cover_page(doc)
    add_front_matter_page(
        doc,
        "Certificate",
        [
            "This is to certify that the project entitled 'Smart Campus Management System (SCMS)' is a bona fide work carried out by the above-mentioned students during Semester IV under the prescribed academic and practical requirements.",
            "The work presented in this report is based on the current implemented system available in the submitted project repository and reflects the design, development, and testing effort completed for academic evaluation and company review.",
            "Guide Signature: ____________________    Date: ____________________",
        ],
    )
    add_front_matter_page(
        doc,
        "Declaration",
        [
            "We declare that this report titled 'Smart Campus Management System (SCMS)' is our original work carried out for Semester IV project submission and internship review purposes.",
            "The report has been prepared using the actual source code, schema, templates, and implemented workflows present in the project workspace. Any external references used for technology understanding have been acknowledged in the references section.",
        ],
    )
    add_front_matter_page(
        doc,
        "Acknowledgement",
        [
            "We express sincere gratitude to our faculty guide, institution, and company mentors for their guidance during the analysis, development, testing, and documentation of this project.",
            "We also acknowledge the practical value of working on a real implementation-oriented codebase, which allowed us to move beyond conceptual design into the engineering details of security, persistence, user workflows, and platform governance.",
        ],
    )
    add_front_matter_page(
        doc,
        "Abstract",
        [
            "Smart Campus Management System (SCMS) is a web-based platform developed to integrate key academic and administrative activities into one structured application. The system provides dedicated portals for students, teachers, college administrators, placement officers, and a platform-level super admin.",
            "The application is implemented using Java 17, Spring Boot 3.2.3, Thymeleaf, MySQL, Spring Data JPA, and supporting libraries such as Apache POI and PDFBox. It covers registration, onboarding, course and class management, attendance, assignments, study materials, examinations, fees, events, notifications, placements, and multi-college administrative governance.",
            "This report documents the project in formal submission format, including problem definition, objectives, system study, design approach, database structure, screen layout analysis, coding strategy, testing methodology, limitations, and future enhancements.",
        ],
    )

    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    add_toc_field(p)
    doc.add_page_break()

    add_intro_sections(doc)
    add_system_study(doc, diagrams)
    add_system_design(doc, diagrams, tables, screens)
    add_coding_section(doc, controllers, models)
    add_testing_section(doc)
    add_closeout_sections(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_report()
